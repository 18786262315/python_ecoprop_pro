
# import imp
import jinja2,pdfkit,base64,time,re,os,requests,json,datetime

from docx import Document
from docx.shared import Inches,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from typing import Optional
from pydantic import BaseModel
# from logging import exception
from ast import literal_eval
from urllib import parse
# from msilib.schema import Error
from fastapi import APIRouter,Form,HTTPException,Body
from .comm import MakeReportlab,getAPI,getDatetimes
from comm.logger import logger
from config import Config
# import ast
# reportlab
from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter, A4
# from reportlab.graphics.charts.piecharts import Pie
from reportlab.lib.units import mm,cm,inch
from reportlab.lib import colors
from reportlab.lib.colors import HexColor,black,red,PCMYKColor
# from reportlab.graphics.shapes import Drawing,Rect
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
# from reportlab.graphics.charts.barcharts import HorizontalBarChart
# from reportlab.lib.formatters import DecimalFormatter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,PageTemplate
from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle
# from hashlib import md5
# from typing import Dict, List, Set, Tuple
# from starlette.responses import FileResponse
from .pro_info_page import XHORIZON_APP_PRO_PDF
from .comparison_pro import XHORIZON_Comparison_PDF
import subprocess

router = APIRouter(prefix="/project",tags=['project'],responses={405: {"description": "Not found"}},)
gettime = getDatetimes()


class Person(BaseModel):
    docPath: Optional[str] = ""
    filepath: Optional[str] = ""
    data: Optional[dict] = {}




@router.get("/")
def read_users():
    return "Mixgo Make PDF API !"


@router.post("/Make_Pdf")
def Make_Pdf(MakeType: str=Form(...),MakeData: str=Form(...)):
    # 创建一个统一接口、按参数区分要生成的PDF,改动较大.暂时没有启用
    logger.info('Make PDF ========>{},Data:{}'.format(MakeType,MakeData))
    
    try:
        MakeData = json.loads(base64.b16decode(MakeData)) # 解码json串
        if MakeType == 'pnd_pro_pdf':
            rentunpath = MakePDF(MakeData.agentId,MakeData.projectId)
        elif MakeType == 'pnd_pro_pdf_Comparison':
            rentunpath = ComparisonPDF(MakeData.agentId,MakeData.projectId)
        elif MakeType == 'xhoapp_pro_pdf':
            rentunpath = XHORIZON_APP_PRO_PDF(MakeData.agentId,MakeData.projectId)
        elif MakeType == 'xhoapp_pro_pdf_Comparison':
            rentunpath = XHORIZON_Comparison_PDF(MakeData.agentId,MakeData.projectId)
        elif MakeType == 'era_bedroom_pdf':
            rentunpath = ERABedroomRports(MakeData.agentId,MakeData.brokeId,MakeData.minPrice,MakeData.maxPrice,MakeData.projectArea,
                                            MakeData.token,MakeData.source)
        elif MakeType == 'ecoprop_shera_pro_pdf':
            rentunpath = Shera_to_Pdf(MakeData.agentId,MakeData.projectId)
        elif MakeType == 'ecoprop_shera_unit_pdf':
            rentunpath = Share_Unit_Pdf(MakeData.agentId,MakeData.unitId)
        elif MakeType == 'ecoprop_shera_pro_cmpare_pdf':
            rentunpath = Share_Pro_compare_Pdf(MakeData.agentId,MakeData.unitId)

        elif MakeType == 'Download_Re_Itinerary_to_Pdf':
            rentunpath = Re_Itinerary(MakeData.customerId)
        elif MakeType == 'Download_Re_Condition_Report_to_Pdf':
            rentunpath = Condition_Report(MakeData.customerId)
        else:
            return {'code':'-1','msg':'error',"datas":'MakeType Error!'}
        return { 'code':'0', 'msg':'succeed', "datas":rentunpath }
    except BaseException as e:
        return {'code':'-1','msg':'error',"datas":e}

def Make_Pdf2(MakeType: str=Form(...),MakeData: str=Form(...)):
    # 创建一个统一接口、按参数区分要生成的PDF,改动较大.暂时没有启用

    pdf_generators ={
        'pnd_pro_pdf': MakePDF,
        'pnd_pro_pdf_Comparison': ComparisonPDF,
        'xhoapp_pro_pdf': XHORIZON_APP_PRO_PDF,
        'xhoapp_pro_pdf_Comparison': XHORIZON_Comparison_PDF,
        'era_bedroom_pdf': ERABedroomRports,
        'ecoprop_shera_pro_pdf': Shera_to_Pdf,
        'ecoprop_shera_unit_pdf': Share_Unit_Pdf,
        'ecoprop_shera_pro_cmpare_pdf': Share_Pro_compare_Pdf,
        'Download_Re_Itinerary_to_Pdf': Re_Itinerary,
        'Download_Re_Condition_Report_to_Pdf': Condition_Report,
    }
    # 验证 MakeType 的合法性
    if MakeType not in pdf_generators:
        logger.error(f"Unsupported MakeType: {MakeType}")
        return {'code': '-1', 'msg': 'error', 'datas': 'MakeType Error!'}
    
    # 解析参数
    try:
        make_data_decoded = json.loads(base64.b16decode(MakeData))
    except json.JSONDecodeError:
        logger.error("JSON decoding error for MakeData")
        return {'code': '-1', 'msg': 'error', 'datas': 'Invalid MakeData format.'}
    
    logger.info('Make PDF ========>{},Data:{}'.format(MakeType,make_data_decoded))
    try:
        # 使用映射的函数生成 PDF,
        # 这里的函数接收参数顺序可能会有不同、所以需要根据函数的参数顺序来调整参数顺序
        generator_function = pdf_generators[MakeType]
        return {'code': '0', 'msg': 'succeed', 'datas': generator_function(**make_data_decoded)}
    except Exception as e:
        logger.error(f"Error generating PDF for type {MakeType}: {e}")
        return {'code': '-1', 'msg': 'error', 'datas': 'An error occurred while generating the PDF.'}


@router.post('/pnd_pro_pdf')
async def GetPndPdfPro(agentId: str=Form(...) ,projectId:str=Form(...)):
    """
    项目ID
    用户ID
    """
    logger.info('创建PDF,%s'%projectId)
    if not agentId and not projectId :
        raise HTTPException(status_code=404, detail="参数错误")

    try:
        rentunpath = MakePDF(agentId,projectId)
        logger.info('PDF创建成功=======>{}'.format(rentunpath))
        rtdata = {
            'code':'0',
            'msg':'succeed',
            "datas":rentunpath
        }
        return rtdata
    except BaseException as e:
            rtdata = {
            'code':'-1',
            'msg':'error',
            "datas":e
            }
            return rtdata

@router.post('/pnd_pro_pdf_Comparison')
async def GetPndComparison(agentId: str=Form(...) ,projectId:str=Form(...)):
    """
    项目ID
    用户ID
    """
    # print(agentId,projectId)
    logger.info('创建PDF,%s'%projectId)
    if not agentId or not projectId :
        raise HTTPException(status_code=404, detail="参数错误")

    try:
        rentunpath = ComparisonPDF(agentId,projectId)
        logger.info('ComparisonPDF创建成功=======>{}'.format(rentunpath))
        rtdata = {
            'code':'0',
            'msg':'succeed',
            "datas":rentunpath
        }
        return rtdata
    except BaseException as e:
            rtdata = {
            'code':'-1',
            'msg':'error',
            "datas":e
            }
            return rtdata

@router.post('/xhoapp_pro_pdf')
async def GetxhoappPdfPro(agentId: str=Form(...) ,projectId:str=Form(...)):
    """
    项目ID
    用户ID
    """
    logger.info('创建PDF,%s'%projectId)
    if not agentId and not projectId :
        raise HTTPException(status_code=404, detail="参数错误")

    try:
        rentunpath = XHORIZON_APP_PRO_PDF(agentId,projectId)
        logger.info('PDF创建成功=======>{}'.format(rentunpath))
        rtdata = {
            'code':'0',
            'msg':'succeed',
            "datas":rentunpath
        }
        return rtdata
    except BaseException as e:
            rtdata = {
            'code':'-1',
            'msg':'error',
            "datas":e
            }
            return rtdata

@router.post('/xhoapp_pro_pdf_Comparison')
async def GetxhoappComparison(agentId: str=Form(...) ,projectId:str=Form(...)):
    """
    项目ID
    用户ID
    """
    # print(agentId,projectId)
    logger.info('创建PDF,%s'%projectId)
    if not agentId or not projectId :
        raise HTTPException(status_code=404, detail="参数错误")
    try:
        rentunpath = XHORIZON_Comparison_PDF(agentId,projectId)
        logger.info('ComparisonPDF创建成功=======>{}'.format(rentunpath))
        rtdata = {
            'code':'0',
            'msg':'succeed',
            "datas":rentunpath
        }
        return rtdata
    except BaseException as e:
            rtdata = {
            'code':'-1',
            'msg':'error',
            "datas":e
            }
            return rtdata

@router.post('/era_bedroom_pdf')
async def GeteraPdfPro(agentId: str=Form(...),
                brokeId: str=Form(...),
                token: str=Form(...),
                minPrice: str=Form(None),
                maxPrice: str=Form(None),
                projectArea:str=Form(None),
                source:str=Form(None),
                ):

    logger.info('----->>>创建 era_bedroom_pdf')
    if not agentId and not brokeId :
        raise HTTPException(status_code=404, detail="参数错误")
    try:
        rentunpath = ERABedroomRports(agentId,brokeId,minPrice,maxPrice,projectArea,token,source)
        logger.info('PDF创建成功=======>{}'.format(rentunpath))
        rtdata = {
            'code':'0',
            'msg':'succeed',
            "datas":rentunpath
        }
        return rtdata
    except BaseException as e:
            rtdata = {
            'code':'-1',
            'msg':'error',
            "datas":e
            }
            return rtdata

@router.post('/ecoprop_shera_pro_pdf')
async def EcopropSheraProPdf(agentId: str=Form(...) ,projectId:str=Form(...)):
    # 生成Ecoprop shera pro pdf
    logger.info('----->>>创建 ecoprop_shera_pro_pdf,agentId:{0},ProjectId:{1}'.format(agentId,projectId))
    if not agentId or not projectId:
        raise HTTPException(status_code=404, detail="参数错误")
    try:
        rentunpath = Shera_to_Pdf(agentId,projectId)
        logger.info('PDF创建成功=======>{}'.format(rentunpath))
        rtdata = {
            'code':'0',
            'msg':'succeed',
            "datas":rentunpath
        }
        return rtdata
    except BaseException as e:
            logger.info('PDF创建异常=======>{}'.format(e))
            rtdata = {
            'code':'-1',
            'msg':'error',
            "datas":e
            }
            return rtdata

@router.post('/ecoprop_shera_unit_pdf')
async def EcopropSheraUnitPdf(agentId: str=Form(...) ,unitId:str=Form(...)):
    # ecoprop pro 单位 pdf
    logger.info('----->>>创建 ecoprop_shera_unit_pdf'+agentId+unitId)
    if not agentId or not unitId:
        raise HTTPException(status_code=404, detail="参数错误")

    try:
        rentunpath = Share_Unit_Pdf(agentId,unitId)
        logger.info('PDF创建成功=======>{}'.format(rentunpath))
        rtdata = {
            'code':'0',
            'msg':'succeed',
            "datas":rentunpath
        }
        return rtdata
    except BaseException as e:
            logger.info('PDF创建异常=======>{}'.format(e))
            rtdata = {
            'code':'-1',
            'msg':'error',
            "datas":e
            }
            return rtdata

@router.post('/ecoprop_shera_pro_cmpare_pdf')
async def EcopropSheraProComparePdf(agentId: str=Form(...) ,projectId:str=Form(...)):

    logger.info('----->>>创建 ecoprop_shera_unit_pdf'+agentId+projectId)
    if not agentId or not projectId:
        raise HTTPException(status_code=404, detail="参数错误")

    try:
        rentunpath = Share_Pro_compare_Pdf(agentId,projectId)
        logger.info('PDF创建成功=======>{}'.format(rentunpath))
        rtdata = {
            'code':'0',
            'msg':'succeed',
            "datas":rentunpath
        }
        return rtdata
    except BaseException as e:
            logger.info('PDF创建异常=======>{}'.format(e))
            rtdata = {
            'code':'-1',
            'msg':'error',
            "datas":e
            }
            return rtdata

@router.post('/Download_Re_Itinerary_to_Pdf')
async def Re_Itinerary(customerId: str=Form(...)):
    """
    Re Lo Sg 项目 
    行程导出PDF功能
    参数：用户ID
    未启用
    """
    logger.info('创建行程导出PDF----->>>{}'.format(customerId))

    getapi = getAPI()
    # host_url = "https://api.singmap.com/relo-api"
    
    try:
        InfoData = {}

        # /customer/queryCustomerItineraryList #父级内容
        ItineraryList = getapi.requsetAPI_POST(Config.ReLoSG_HOST+"/customer/queryCustomerItineraryList",
                                                params={"pageNo": "1",
                                                        "pageSize": "1000",
                                                        "customerId": customerId,
                                                        "requestTime": int(time.time()* 1000)},
                                                item='ReLoSG')
        
        # /customer/queryCustomerItineraryItem #子级内容
        InfoData['ItineraryDataList'] = []
        for ItineraryParent in ItineraryList['lists']:
            ItineraryDataItem = {"itineraryId": ItineraryParent['itineraryId'],"requestTime": int(time.time()* 1000)}
            ItineraryData = getapi.requsetAPI_POST(Config.ReLoSG_HOST+"/customer/queryCustomerItineraryItem",params=ItineraryDataItem,item='ReLoSG')
            if ItineraryData == []: continue
            # 转换拼接父级和子级内容,并整合到一个列表
            InfoData['ItineraryDataList'].extend(list(map(lambda item: {**item, **ItineraryParent}, ItineraryData))) 

        # /customer/info # 客户详细信息
        InfoData['UserInfo'] = getapi.requsetAPI_POST(Config.ReLoSG_HOST+"/customer/info",
                                            params={"customerId": customerId, "requestTime": int(time.time()* 1000)},
                                            item='ReLoSG')

        # 列表按 _itineraryDate 和 startDate进行倒序排列
        InfoData['ItineraryDataList'] = sorted(InfoData['ItineraryDataList'], key=lambda x: (x['_itineraryDate'], x['startDate']), reverse=True)
        InfoData['NowDate'] = gettime.getDate()

        # 公司列表
        #  https://api.singmap.com/relo-api/company/list

        companyList = getapi.requsetAPI(Config.ReLoSG_HOST+"/company/list",
                                                params={
                                                        "name": '',
                                                        "pageNo": "1",
                                                        "pageSize": "1000",
                                                        "requestTime": int(time.time()* 1000)},
                                                item='ReLoSG')
        # 取值公司名称
        condition = lambda x: x['companyId'] == InfoData['UserInfo']['companyId']
        InfoData['UserInfo']['companyName'] = next(filter(condition, companyList['lists']), '')['companyName']

        # 管理员列表
        # https://api.singmap.com/relo-api/user/list
        userList = getapi.requsetAPI(Config.ReLoSG_HOST+"/user/list",
                                                params={"pageNo": "1",
                                                        "pageSize": "10000",
                                                        "status": '',
                                                        "userName": '',
                                                        "role": '',
                                                        "requestTime": int(time.time()* 1000)},
                                                item='ReLoSG')

        # 取值管理员名称
        condition1 = lambda x: x['userId'] == InfoData['UserInfo']['mangerUserId']
        mangerUserInfo = next(filter(condition1, userList['lists']), '')
        InfoData['UserInfo']['HandlerName'] = mangerUserInfo['firstName'] +'' + mangerUserInfo['lastName']
        
        
        # 输出文件路径
        new_file_name = InfoData['UserInfo']['firstName']+' '+InfoData['UserInfo']['lastName']+"-Schedule"+'.pdf'
        re_path = os.path.join(Config.ecoprop_return_path,'re_itinerary',InfoData['NowDate'],new_file_name).replace('\\','/')
        logger.info('Set return Path ====>>>>'+re_path)
        # 文件夹检查
        if not os.path.exists(os.path.split(re_path)[0]): 
            os.makedirs(os.path.split(re_path)[0])

        logger.info('Get Temp ====>>>> re_Schedule.html')
        env = jinja2.Environment(loader=jinja2.FileSystemLoader(searchpath=Config.ecoprop_temp_path,encoding='utf-8'))
        template = env.get_template('re_Schedule.html')

        # 模板填充参数
        options = {
                # 'page-width': '842px',
                # 'page-height': '595px',
                'page-size': ['A4'],
                'margin-top': '0mm',
                'margin-right': '0mm',
                'margin-bottom': '0mm',
                'margin-left': '0mm',
                'orientation':'Landscape', #横向
                'encoding': "UTF-8",
                'no-outline': None,

                'enable-local-file-access': None,  # 允许访问本地文件
                'javascript-delay': 2000,          # 延迟执行JavaScript，确保图片完全加载
                'images': None,                    # 启用图像加载
                'enable-javascript': None,         # 启用JavaScript

                # 日志设置
                'quiet': None,                     # 减少输出信息

                # 资源加载控制
                # 内存和超时设置
                'disable-smart-shrinking': None,   # 禁用智能缩小，保持原始尺寸比例
                'zoom': 1,                         # 缩放因子
                'load-error-handling': 'ignore',   # 忽略加载错误
                'load-media-error-handling': 'ignore',  # 忽略媒体加载错误
            }
        try:
            logger.info('Add tmp Info ====>>>>')
            InfoData = eval(re.sub('None','\'\'',str(InfoData))) # 去除None值
            htmls = template.render(InfoData)
            config = pdfkit_config()

            
            pdfkit.from_string(htmls,re_path,options=options,configuration=config)
        except Exception as e:
            logger.error(e)
            raise HTTPException(status_code=404, detail="PDF Error")
        logger.info('PDF创建成功=======>{}'.format(re_path))
        # rtdata = {'code':'0','msg':'succeed',"datas":re_path}
        return {'code':'0','msg':'succeed',"datas":re_path}
    except BaseException as e:
        logger.info('PDF创建异常=======>{}'.format(e))
        # rtdata = {'code':'-1','msg':'error',"datas":e}
        return {'code':'-1','msg':'error',"datas":e}


@router.post('/Download_Re_Condition_Report_to_Pdf')
async def Condition_Report(customerId: str=Form(...)):
    """
    Re Lo Sg 项目
    反馈信息导出PDF功能
    参数：客户ID
    """
    logger.info('客户反馈信息导出PDF----->>>{}'.format(customerId))
    
    try:
        getapi = getAPI()
        Condition_Data = {}
        # /customer/info # 客户详细信息
        Condition_Data['UserInfo'] = getapi.requsetAPI_POST(Config.ReLoSG_HOST+"/customer/info",
                                            params={
                                                    "customerId": customerId,
                                                    "requestTime": int(time.time()* 1000)
                                                    },
                                            item='ReLoSG')
        
        # /customer/queryCustomerFeedbackList # 获取反馈信息
        ConditionList = getapi.requsetAPI(Config.ReLoSG_HOST+"/customer/queryCustomerFeedbackList",
                                        params={
                                                "customerId": customerId,
                                                "pageNo": "1",
                                                "pageSize": "1000",
                                                "requestTime": int(time.time()* 1000)
                                                },
                                        item='ReLoSG')
        Condition_Data['Conditionlist'] = ConditionList['lists']

        # 多个图片分割处理
        for index,data in enumerate(Condition_Data['Conditionlist']):
            Condition_Data['Conditionlist'][index]['photo'] = data['photo'].split(',')

        # 输出文件路径
        new_file_name ="Condition_Report_"+Condition_Data['UserInfo']['firstName']+' '+Condition_Data['UserInfo']['lastName']+'.pdf'
        rt_path = os.path.join(Config.ecoprop_return_path,'re_itinerary',gettime.getDate(),new_file_name).replace('\\','/')

        # 文件夹检查
        if not os.path.exists(os.path.split(rt_path)[0]): 
            os.makedirs(os.path.split(rt_path)[0])

        logger.info('Jinja Temp Filling ====>>>> Condition_Report.html')
        env = jinja2.Environment(loader=jinja2.FileSystemLoader(searchpath=Config.ecoprop_temp_path,encoding='utf-8'))
        template = env.get_template('Condition_Report.html')

        # 模板填充参数
        options = {
            'page-size': ['A4'],
            'margin-top': '15mm',
            'margin-right': '5mm',
            'margin-bottom': '10mm',
            'margin-left': '5mm',
            'orientation':'portrait',
            'encoding': "UTF-8",
            # 'no-outline': None,
            # 'header-html':Config.ecoprop_temp_path+'/pdfHeader.html', #设置页眉数据，作为页眉的html页面必须有<!DOCTYPE html> 不能动态化
            # '--header-center':'Condition Report',
            '--header-left':'Condition Report',
            '--header-line':'--header-line',
            '--header-spacing':5,
            '--footer-right':'[page]/[topage]',
            # '--footer-line':'--footer-line',
            '--footer-spacing':5,
        }
        try:
            logger.info('Add tmp Info ====>>>>')

            # 去除None值
            datas = eval(re.sub('None','\'\'',str(Condition_Data))) 
            htmls = template.render(datas)

            # linux指定位置配置，因找不到wkhtmltopdf导致无法生成PDF 问题处理方案
            config = pdfkit_config()

            open(rt_path, 'w').close() 

            pdfkit.from_string(htmls,rt_path,options=options,configuration=config)
        except Exception as e:
            logger.error('PDF Create Error=======>{}'.format(e))
            raise HTTPException(status_code=404, detail="PDF Error")
        logger.info('PDF Create Overall=======>{}'.format(rt_path))
        return {'code':'0','msg':'succeed',"datas":rt_path}
    except BaseException as e:
        logger.error('PDF Create Error=======>{}'.format(e))
        return {'code':'-1','msg':'error',"datas":e}

# PND项目PDF生成函数
def MakePDF(agentId,projectId):
    # 基础数据准备 =====================================================
    getapi = getAPI()
    # 项目信息
    proinfourl = Config.API_IP+Config.PND_PROJECT_INFO
    proinfodata = {
        "projectId":projectId,
        "agentId":agentId
    }
    logger.info('项目信息查询')
    
    proinfo = getapi.requsetAPI(proinfourl,proinfodata)
    if "projectInfo" in proinfo and "unitInfo" in proinfo and "agentInfo" in proinfo:
        logger.info('项目信息查询成功--->>>projectId:%s,agentId:%s,projectInfo:%s,unitInfo:%s,agentInfo:%s'%(projectId,agentId,proinfo['projectInfo'],proinfo['unitInfo'],proinfo['agentInfo']))
    else:
        logger.info('项目信息查询失败--->>>projectId:%s,agentId:%s'%(projectId,agentId))
        raise HTTPException(status_code=404, detail="项目信息查询失败")
    
    prodatainfo = proinfo['projectInfo'] # 项目信息
    unitInfo = proinfo['unitInfo'] # 单位信息
    userinfo = proinfo['agentInfo'] #用户信息
    dealInfo = proinfo['dealInfo'] # 房间-信息
    Symbol = "$" #价格符
    if prodatainfo["currencySymbol"] != None:
        Symbol = prodatainfo["currencySymbol"]
    unitproce = [0,0,0,0,0] #单位售价

    # 项目PDF上传图片
    # propdfurls = Config.API_IP+Config.PND_PDF_PROJECT_LIST
    # propdfdatas = {
    #     "projectId":projectId,
    #     "type":""
    # }
    propdfinfo = getapi.requsetAPI( Config.API_IP+Config.PND_PDF_PROJECT_LIST,{
        "projectId":projectId,
        "type":""
    })

    # if propdfinfo == None :
    #     logger.info('项目PDF上传图片为空--->>>projectId:%s,agentId:%s'%(projectId,agentId))

    # 项目区域
    # districturl = Config.API_IP+Config.PND_PDF_DISTRICT_LIST
    # districtdata = {
    #     "district":prodatainfo['district'],
    #     "type":"1"
    # }
    districtinfo = getapi.requsetAPI(Config.API_IP+Config.PND_PDF_DISTRICT_LIST,{
        "district":prodatainfo['district'],
        "type":"1"
    })
    logger.info('项目区域销售图查询成功--->>>%s'%(districtinfo))

    # PDF文件　
    # fileurl = Config.API_IP+"/pnd-api/pdf/queryPdfList"
    # filedata = {
    #     "fileName":"ProjectReport",
    #     "page":""
    # }
    fileinfo = getapi.requsetAPI(Config.API_IP+"/pnd-api/pdf/queryPdfList",{
        "fileName":"ProjectReport",
        "page":""
    })
    logger.info('PDF文件页面图片查询成功--->>>%s'%(fileinfo))

    #新加坡区间 销售统计
    regionurl = Config.API_IP+Config.PND_PROJECT_RETAIL_COUNT
    logger.info('API--->>>%s'%(regionurl))
    # RCRdata = { "region":"RCR"}
    RCRinfo = getapi.requsetAPI(regionurl,{ "region":"RCR"})
    logger.info('RCR查询成功--->>>%s'%(RCRinfo))

    # CCRdata = { "region":"CCR"}
    CCRinfo = getapi.requsetAPI(regionurl,{ "region":"CCR"})
    logger.info('CCR查询成功--->>>%s'%(CCRinfo))

    # OCRdata = {"region":"OCR"}
    OCRinfo = getapi.requsetAPI(regionurl,{"region":"OCR"})
    logger.info('OCR查询成功--->>>%s'%(OCRinfo))

    # 创建PDF文档 =====================================================
    song = "simsun"

    # 注册字体    
    msyh = os.path.join(os.getcwd(), "font", "msyh.ttf")
    msyhbd = os.path.join(os.getcwd(), "font", "msyhbd.ttf")
    ARIALBD = os.path.join(os.getcwd(), "font", "ARIALBD.TTF")
    arial = os.path.join(os.getcwd(), "font", "arial.ttf")
    simsun = os.path.join(os.getcwd(), "font", "simsun.ttc")

    pdfmetrics.registerFont(TTFont(song, simsun))
    pdfmetrics.registerFont(TTFont('ARIALBD',ARIALBD)) #注册字体
    pdfmetrics.registerFont(TTFont('arial',arial)) #注册字体
    pdfmetrics.registerFont(TTFont('msyh',msyh)) #注册字体
    pdfmetrics.registerFont(TTFont('msyhbd',msyhbd)) #注册字体
    Imagepath = os.path.join(Config.filepath,'file')

    pagesize = (1747,965) # 画布大小
    # pagesize = (A4[1],A4[0]) # 画布大小
    # PND文件夹+ 项目ID + 用户信息 + 文件名称
    gettime = getDatetimes()
    tt = gettime.getDate()
    uppath = os.path.join(Config.returnpaths,tt)

    if not os.path.exists(uppath):
        os.makedirs(uppath)
    savepath = os.path.join(uppath,str(int(time.time()))+'.pdf') 
    logger.info(f'---------->>>文档创建{savepath}')

    # returnPath = os.path.join(Config.filepath,tt,str(int(time.time()))+'.pdf')
    doc = canvas.Canvas(savepath,pagesize=pagesize)
    doc.setTitle(prodatainfo['projectName'])

    makefunc = MakeReportlab(doc,Imagepath,pagesize,Symbol) # 加载方法

    ######################################################################
    # page1  ===============================================
    # 背景
    makefunc.background('0.jpg')
    if userinfo['logo'] and Config.envs == 'release':

        makefunc.AddURLImages(Config.imgpath+userinfo['logo'],x=650,y=150,w=224,h=224) 

    agentdata4 = [
    ["NAME",':', userinfo['agentName']],
    ["MOBILE",':', userinfo['mobile']],
    ["EMAIL",':',userinfo['email']],
    ["CEA",':', userinfo['regNum']]
    ]

    t = Table(agentdata4, style={
    ("FONT", (0, 0), (-1, -1), 'arial', 35,50),
    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
    ('ALIGN', (1, 0), (1, -1), 'CENTER')
    })
    t._argW[1] = 20
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 880, 150)
    doc.showPage()  # 保存当前画布页面
    logger.info('============>> 1 page')

    # Page2  ===============================================
    makefunc.AddImages('1.jpg',w=pagesize[0],h=pagesize[1])
    doc.showPage()
    logger.info('============>> 2 page')

    # Page3 ===============================================
    makefunc.background('bgB.png')
    my_text = prodatainfo['projectName']
    doc.setFillColorRGB(0,0,0) #choose your font colour
    doc.setFont('msyhbd', 60)
    textobject = doc.beginText(70, pagesize[1]-80)
    for line in my_text.splitlines(False):
        textobject.textLine(line.rstrip())
    doc.drawText(textobject)

    makefunc.addTesxts(fontsize=25,fontname='ARIALBD',x=70,y=pagesize[1]-150,text="PROJECT SUMMARY",color=HexColor('#E37200'))
    # 项目主图
    if prodatainfo['mainImage'] and Config.envs != "test":
        ...
        makefunc.ImageAdaptive(Config.imgpath+prodatainfo['mainImage'],x=pagesize[0]-400,y=pagesize[1]-400,w=400,h=200) 
    completionDate = ''
    if type(prodatainfo['completionDate']) is int:
        timeArray = time.strptime(prodatainfo['completionDate'], "%Y-%m-%d")
        completionDate = time.strftime("%Y", timeArray)
    if type(prodatainfo['completionDate']) is str and makefunc.isVaildDate(prodatainfo['completionDate']):
        a = time.strptime(prodatainfo['completionDate'], "%Y-%m-%d %H:%M:%S")
        completionDate = time.strftime("%Y", a)
    else:
        completionDate = prodatainfo['completionDate']
    prodata = [
    ["Developer",':', prodatainfo['brokeName']],
    ["Tenure",':', prodatainfo['tenure']],
    ["District",':', prodatainfo['district']],
    ["Region",':', prodatainfo['projectArea']],
    ["Top",':', completionDate],
    ["Total Units",':', prodatainfo['unitsNum']]
    ]

    makefunc.drawUserInfoTable(prodata,70,580)
    logger.info('============>> Pro info OK !')

    # 单位类型价格统计表
    # makefunc.addTesxts(fontsize=25,fontname='ARIALBD',x=70,y=550,text="UNIT PRICE",color=HexColor('#E37200'))
    prodataroombed = [
    ['UNIT PRICE'],
    ["Types",'','Size Range','Price From'],
    ]
    Pielist = []
    for item in unitInfo:
        logger.info(item)
        size = '-'
        if item['minArea'] != 0 and item['maxArea'] != 0: # 户型最大最小面积
            size = '{0}-{1}sqft'.format(item['minArea'],item['maxArea'])

        if item['bedrooms'] == None or item['bedrooms'] < 1 or item['bedrooms'] > 5:
            continue
        elif item['price']:
            prodataroombed.append([str(item['bedrooms'])+" Bedroom",':',size,makefunc.priceset(item['price'])])
            # 
            if item['bedrooms'] == 1:
                unitproce[0] = item['price']
            if item['bedrooms'] == 2:
                unitproce[1] = item['price']
            if item['bedrooms'] == 3:
                unitproce[2] = item['price']
            if item['bedrooms'] == 4:
                unitproce[3] = item['price']
            if item['bedrooms'] == 5:
                unitproce[4] = item['price']
        else:
            prodataroombed.append([str(item['bedrooms'])+" Bedroom",':',size,makefunc.priceset(item['price'])])
            # if item['bedrooms'] == 1:
            #     unitproce[0] = 0
            # if item['bedrooms'] == 2:
            #     unitproce[1] = 0
            # if item['bedrooms'] == 3:
            #     unitproce[2] = 0
            # if item['bedrooms'] == 4:
            #     unitproce[3] = 0
            # if item['bedrooms'] == 5:
            #     unitproce[4] = 0
    # makefunc.drawUserInfoTable(prodataroombed,70,300)

    t = Table(prodataroombed, style={
    ("FONT", (0, 0), (-1, 1), 'ARIALBD',25),
    # ("FONT", (0, 1), (-1, 1), 'ARIALBD',25),
    ("TEXTCOLOR", (0, 0), (-1, 0), HexColor('#E37200')),
    ("FONT", (0, 1), (-1, -1), 'arial', 26),
    ("TEXTCOLOR", (0, 1), (-1, -1), colors.black),
    ('ALIGN', (1, 0), (1, -1), 'CENTER')
    })
    t._argW[1] = 20
    t._argW[2] = 250
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 70, 300)
    logger.info('============>> Pro unit Info OK !')

    makefunc.addTesxts(fontsize=25,fontname='ARIALBD',x=pagesize[0]-930,y=pagesize[1]-250,text="BEDROOM UNITS SHARES",color=HexColor('#E37200'))
    # 饼图
    makefunc.MakePie_PND(unitInfo).drawOn(doc,830,250)
    logger.info('============>> Pro MakePie OK !')

    # 底部内容
    page3_btm = [
        ['Project Brochure','[-]'],
        ['360 Panorama','[-]'],
        ['School(s) Within 1 KM','[-]'],
        ['Nearby MRT within 2KM','[-]'],
    ]
    if prodatainfo['url'] != None:
        page3_btm[0][1] = 'CLICK HERE'
        doc.linkURL(Config.imgpath+prodatainfo['url'], (280,240,280+150,240+22))
    if prodatainfo['ivtList'] != []:
        page3_btm[1][1] = 'CLICK HERE'
        # print(prodatainfo['ivtList'][0])
        doc.linkURL(prodatainfo['ivtList'][0], (280,215,280+150,215+22))
    if prodatainfo['facilitiesMap']:
        facilities = json.loads(prodatainfo['facilitiesMap'])
        for item in facilities:
            if item['type'] == 'subway_station' and item['value']:
                # print(item)
                m =[re.sub('MRT Station|Station','',MRT['name']) for MRT in item['value']] #多个地铁
                # for MRT in item['value']:
                    # m.append(re.sub('MRT Station|Station','',MRT['name']))
                # print('/'.join(m))
                page3_btm[3][1] = makefunc.create_body_text('/'.join(m),color=colors.black) 
            if item['type'] == 'school' and item['value']:
                school =[MRT['name'] for MRT in item['value']]
                page3_btm[2][1] = makefunc.create_body_text('/'.join(school) if len(school)<3 else '/'.join(school[0:3]),color=colors.black)
                # page3_btm[0][1] = item['value'][0]['name']
                # for school in item['value']:
                #     print(school['name'])
    t = Table(page3_btm, style={
    ("FONT", (0, 0), (-1, -1), "msyhbd", 20),
    ("TEXTCOLOR", (0, 0), (0, -1), HexColor('#E37200')),
    })
    t._argW[1] = 900
    t._argH[2] = 50 
    t._argH[3] = 60 
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 70, 100)
    # 项目 IVT跳转\楼书跳转
    logger.info('============>> Pro Bottom OK !')

    makefunc.addTesxts(fontsize=16,x=10,y=25,text="Source: Google Maps")
    # 页脚
    makefunc.addTesxts(fontsize=16,x=10,y=10,text="Personalised Property Analytics Report • Value-adding Your Home Purchase")
    doc.showPage()  # 保存当前画布页面
    logger.info('============>> PROJECT SUMMARY')

    # Page4 项目周边===========================================================================
    # 背景图
    makefunc.background('bgB.png')
    makefunc.addTesxts(fontsize=60,x=70,y=pagesize[1]-80,text=prodatainfo['projectName'],fontname='msyhbd')
    makefunc.addTesxts(fontsize=40,x=pagesize[0]/2.6,y=pagesize[1]-270,text="LOCATION HIGHLIGHT",color=HexColor('#E37200'))
    if prodatainfo['snapshotLogo'] and Config.envs != "test":
        # 地址Google 截图
        makefunc.AddURLImages(Config.imgpath+prodatainfo['snapshotLogo'],x=100,y=100,w=700,h=550) 
    if propdfinfo != [] :
        # 周边截图
        makefunc.AddURLImages(Config.imgpath+propdfinfo[0]['logo'],x=pagesize[0]/2,y=100,w=700,h=550)
    makefunc.addTesxts(fontsize=18,x=100,y=50,text="Source: Google Maps")
    makefunc.addTesxts(fontsize=18,x=pagesize[0]/2,y=50,text="Source: URA Map")
    doc.showPage()  # 保存当前画布页面
    logger.info('============>> LOCATION HIGHLIGHT')

    # Page6 区域出租记录图片 ===========================================================================
    if districtinfo != [] and Config.envs != "test":

        makefunc.AddURLImages(Config.imgpath+districtinfo[0]['logo'],w=pagesize[0],h=pagesize[1])
        makefunc.addTesxts(fontsize=50,x=10,y=pagesize[1]-80,text='District Pricing (Rental)',color=HexColor('#E37200'))
        doc.showPage()  # 保存当前画布页面
        logger.info('============>> District Pricing (Rental)')

    # Page7 中介信息   ===========================================================================
    # makefunc.background('6.jpg')
    # if userinfo['logo'] and Config.envs == 'release' and Config.envs != "test":
    #     makefunc.AddURLImages(Config.imgpath+userinfo['logo'],x=650,y=200,w=224,h=224) 

    # agentdata2 = [
    # ["NAME",':', userinfo['agentName']],
    # ["MOBILE",':', userinfo['mobile']],
    # ["EMAIL",':',userinfo['email']],
    # ["CEA",':', userinfo['regNum']]
    # ]
    # t = Table(agentdata2, style={
    # ("FONT", (0, 0), (-1, -1), "arial", 35,50),
    # ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
    # ('ALIGN', (1, 0), (1, -1), 'CENTER')
    # })
    # t._argW[1] = 20
    # t.wrapOn(doc, 0, 0)
    # t.drawOn(doc, 880, 195)
    # # WhatsApp 聊天跳转按钮
    # doc.linkURL('https://api.whatsapp.com/send?phone= '+userinfo['mobile'], (1250,90,1250+447,150))
    # doc.showPage()  # 保存当前画布页面
    # logger.info('============>> Agent Info OK !')

    # Page8 ===========================================================================
    makefunc.background('bgA.png')
    makefunc.addTesxts(fontsize=40,x=pagesize[0]/3,y=650,text="Guide to Financial Wellness")
    # 单位统计表
    unittype2 = [
    ["Unit Type",'Price from', "Monthly Installment","Min. Monthly Income \n Required For The Purchase",""],
    ["",'','', "Employee","Self Employed"],
    ]
    for item in unitInfo:
        if item['bedrooms'] == None  or item['bedrooms'] < 1 or item['bedrooms'] > 5:
            continue
        elif item['price'] :
            stage = makefunc.stages_numb(item['price'])
            if stage != '-':
                unittype2.append([str(item['bedrooms'])+\
                    " Bedroom",makefunc.priceset(item['price']),\
                        makefunc.stages(item['price']), \
                            makefunc.priceset(makefunc.stages_numb(item['price'])/0.55), \
                                makefunc.priceset((makefunc.stages_numb(item['price'])/0.55)/0.7)])
            else:
                unittype2.append([str(item['bedrooms'])+\
                    " Bedroom",makefunc.priceset(item['price']),\
                        makefunc.stages(item['price']), \
                            '-', \
                                '-'])
        else:
            unittype2.append([str(item['bedrooms'])+" Bedroom",'-','-','-','-'])
    
    t = Table(unittype2,(pagesize[0]-436)/5,50, style={
    ('SPAN', (-2, 0), (-1,0)), # 合并单元格(列,行)
    ('BACKGROUND',(1,2),(-1,-1), colors.white),
    ('LINEBEFORE', (0, 0), (-1, -1), 0.1 * cm, colors.black),
    ('BOX', (0, 0), (-1, -1), 0.1 * cm, colors.black),
    ('LINEABOVE', (0, 0), (-1, -1), 0.1 * cm, colors.black),
    ('LINEBELOW', (0,0), (-1,-1), 0.1 * cm, colors.black),
    # ("FONT", (0, 0), (-1, -1), song, 25),
    ("FONT", (0, 0), (-1, -1), "msyh", 25),
    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER')
    })
    t._argH[0] = 75 
    t._argH[1] = 65 
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 200, 250)

    # 页脚
    # makefunc.addTesxts(fontsize=16,x=70,y=30,text=" PERSONALISED PROPERTY ANALYTICS REPORT. | 2022")
    makefunc.addTesxts(fontsize=16,x=300,y=200,text=R"*Calculation based on 30 years tenure, 75% LTV, 3.5% bank interest rate. For your personal financial calculation, please approach our sales person for assistance.")

    doc.showPage()  # 保存当前画布页面
    logger.info('============>>Guide to Financial Wellness')
    
    # Page9 ===========================================================================
    # 背景
    # makefunc.background('8.jpg')
    # if userinfo['logo'] and Config.envs == 'release':
    #     makefunc.AddURLImages(Config.imgpath+userinfo['logo'],x=100,y=250,w=224,h=224) 

    # agentdata5 = [
    # ["NAME",':', userinfo['agentName']],
    # ["MOBILE",':', userinfo['mobile']],
    # ["EMAIL",':',userinfo['email']],
    # ["CEA",':', userinfo['regNum']]
    # ]

    # t = Table(agentdata5, style={
    # ("FONT", (0, 0), (-1, -1), 'arial', 35,50),
    # ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
    # ('ALIGN', (1, 0), (1, -1), 'CENTER')
    # })
    # t._argW[1] = 20
    # t.wrapOn(doc, 0, 0)
    # t.drawOn(doc, 350, 250)
    # # WhatsApp 聊天跳转按钮
    # doc.linkURL('https://api.whatsapp.com/send?phone='+userinfo['mobile'], (600,120,600+447,196))

    # doc.showPage()  # 保存当前画布页面

    # Page10 ===========================================================================
    makefunc.background('9.jpg')
    makefunc.background('bgA.png')

    makefunc.addTesxts(fontsize=60,x=400,y=pagesize[1]-100,text="PROGRESSIVE PAYMENT")
    # 贷款计算
    unitcalculator = [
        ['','','1 Bedroom','2 Bedroom','3 Bedroom','4 Bedroom','5 Bedroom'],
        ['','','-','-','-','-','-'],
        ['','','-','-','-','-','-'],
        ['','','','','',''],
        ['','','-','-','-','-','-'],
        ['','','-','-','-','-','-'],
        ['','','-','-','-','-','-'],
        ['','','-','-','-','-','-'],
        ['','','','','',''],
        ['','','-','-','-','-','-'],
        ['','','-','-','-','-','-'],
    ]
    for y,data in enumerate(unitcalculator):
        # if item > 1:
        for x,data2 in enumerate(data):
            if x >1:
                BSD = makefunc.priceBSD(unitproce[x-2])
                if y ==1:
                    unitcalculator[y][x] = makefunc.testNan(Symbol,round(unitproce[x-2]))
                if y ==2:
                    unitcalculator[y][x] = makefunc.testNan(Symbol,round(unitproce[x-2]*0.75))
                if y ==4:
                    unitcalculator[y][x] = makefunc.testNan(Symbol,round(unitproce[x-2]*0.05))
                if y ==5:
                    unitcalculator[y][x] = makefunc.testNan(Symbol,BSD)
                if y ==6:
                    unitcalculator[y][x] = makefunc.testNan(Symbol,round(unitproce[x-2]*0.15))
                if y ==7:
                    unitcalculator[y][x] = makefunc.testNan(Symbol,round((unitproce[x-2]*0.2)+BSD))
                if y ==9:
                    unitcalculator[y][x] = makefunc.testNan(Symbol,round(unitproce[x-2]*0.05)) 
                if y ==10:
                    unitcalculator[y][x] = makefunc.testNan(Symbol,round(((unitproce[x-2]*0.2)+BSD)+(unitproce[x-2]*0.05)))
    t = Table(unitcalculator,220,32)
    style={
    ('SPAN', (0, 0), (1,0)), # 合并单元格(列,行)
    ('LINEBEFORE', (0, 0), (-1, -1), 0.05 * cm, colors.whitesmoke),
    ("TEXTCOLOR", (0,1), (-1, -1), '#1F5B73'),
    ('BACKGROUND',(2,1),(-1,2), colors.white),
    ('BACKGROUND',(2,4),(-1,6), colors.white),

    ("FONT", (2,7),(-1,7),'ARIALBD', 16,25),
    ('BACKGROUND',(2,7),(-1,7), colors.orange),
    ("TEXTCOLOR", (2,7),(-1,7), colors.white),

    ('BACKGROUND',(2,9),(-1,9), colors.white),

    ("FONT", (2,10),(-1,10),'ARIALBD', 16,25),
    ('BACKGROUND',(2,10),(-1,10), colors.orange),
    ("TEXTCOLOR", (2,10),(-1,10), colors.white),

    ("FONT", (0, 0), (-1, -1),'arial', 16,25),

    ('ALIGN', (1, 0), (-1, -1), 'CENTER')
    }
    t._argW[0] = 100
    t._argW[1] = 200
    t.setStyle(style)
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 150, 442)

    unitcalculator1 = [
        ['','','Outstanding \n Loan','Monthly \n Installment','Outstanding \n Loan','Monthly \n Installment','Outstanding \n Loan','Monthly \n Installment','Outstanding \n Loan','Monthly \n Installment','Outstanding \n Loan','Monthly \n Installment'],
        ['','',makefunc.counts(Symbol,unitproce[0],0.05),makefunc.counts(Symbol,unitproce[0],0.05,0.0035),
                                            makefunc.counts(Symbol,unitproce[1],0.05),makefunc.counts(Symbol,unitproce[1],0.05,0.0035),
                                            makefunc.counts(Symbol,unitproce[2],0.05),makefunc.counts(Symbol,unitproce[2],0.05,0.0035),
                                            makefunc.counts(Symbol,unitproce[3],0.05),makefunc.counts(Symbol,unitproce[3],0.05,0.0035),
                                            makefunc.counts(Symbol,unitproce[4],0.05),makefunc.counts(Symbol,unitproce[4],0.05,0.0035),],
        ['','',makefunc.counts(Symbol,unitproce[0],0.15),makefunc.counts(Symbol,unitproce[0],0.15,0.0035),
                                            makefunc.counts(Symbol,unitproce[1],0.15),makefunc.counts(Symbol,unitproce[1],0.15,0.0035),
                                            makefunc.counts(Symbol,unitproce[2],0.15),makefunc.counts(Symbol,unitproce[2],0.15,0.0035),
                                            makefunc.counts(Symbol,unitproce[3],0.15),makefunc.counts(Symbol,unitproce[3],0.15,0.0035),
                                            makefunc.counts(Symbol,unitproce[4],0.15),makefunc.counts(Symbol,unitproce[4],0.15,0.0035),],
        ['','',makefunc.counts(Symbol,unitproce[0],0.35),makefunc.counts(Symbol,unitproce[0],0.35,0.0035),
                                                    makefunc.counts(Symbol,unitproce[1],0.35),makefunc.counts(Symbol,unitproce[1],0.35,0.0035),
                                                    makefunc.counts(Symbol,unitproce[2],0.35),makefunc.counts(Symbol,unitproce[2],0.35,0.0035),
                                                    makefunc.counts(Symbol,unitproce[3],0.35),makefunc.counts(Symbol,unitproce[3],0.35,0.0035),
                                                    makefunc.counts(Symbol,unitproce[4],0.35),makefunc.counts(Symbol,unitproce[4],0.35,0.0035),],
        ['','','','',''],
        ['','',makefunc.counts(Symbol,unitproce[0],0.6),makefunc.counts(Symbol,unitproce[0],0.6,0.0035),
                            makefunc.counts(Symbol,unitproce[1],0.6),makefunc.counts(Symbol,unitproce[1],0.6,0.0035),
                            makefunc.counts(Symbol,unitproce[2],0.6),makefunc.counts(Symbol,unitproce[2],0.6,0.0035),
                            makefunc.counts(Symbol,unitproce[3],0.6),makefunc.counts(Symbol,unitproce[3],0.6,0.0035),
                            makefunc.counts(Symbol,unitproce[4],0.6),makefunc.counts(Symbol,unitproce[4],0.6,0.0035),],
        ['','',makefunc.counts(Symbol,unitproce[0],0.75),makefunc.counts(Symbol,unitproce[0],0.75,0.0035),
                    makefunc.counts(Symbol,unitproce[1],0.75),makefunc.counts(Symbol,unitproce[1],0.75,0.0035),
                    makefunc.counts(Symbol,unitproce[2],0.75),makefunc.counts(Symbol,unitproce[2],0.75,0.0035),
                    makefunc.counts(Symbol,unitproce[3],0.75),makefunc.counts(Symbol,unitproce[3],0.75,0.0035),
                    makefunc.counts(Symbol,unitproce[4],0.75),makefunc.counts(Symbol,unitproce[4],0.75,0.0035),],

    ]
    logger.info('============>>Guide to Financial Wellness')
    
    unitcalculatortb1 = Table(unitcalculator1,110,35, style={
    ('SPAN', (0, 0), (1,0)), # 合并单元格(列,行)
    ('LINEBEFORE', (0, 0), (-1, -1), 0.05 * cm, colors.whitesmoke),
    ("FONT", (0, 0), (-1, -1), 'arial', 16),
    ('FONTNAME', (0,0), (-1,0), 'Courier-Bold'),
    ('TEXTCOLOR',(2,0),(-1,0), colors.white),
    ('BACKGROUND',(2,0),(-1,0), colors.midnightblue),
    ("FONT", (2,0),(-1,0), 'arial', 14),

    ('BACKGROUND',(2,1),(-1,3), colors.white),
    ('BACKGROUND',(2,5),(-1,6), colors.white),
    ("TEXTCOLOR", (0,1), (-1, -1), '#1F5B73'),
    ('ALIGN', (1, 0), (-1, -1), 'CENTER')
    })
    unitcalculatortb1._argH[0] = 50
    unitcalculatortb1._argW[0] = 100
    unitcalculatortb1._argW[1] = 200
    unitcalculatortb1.wrapOn(doc, 0, 0)
    unitcalculatortb1.drawOn(doc, 150, 160)

    makefunc.AddImages('left_report.png',x=150,y=155,w=280,h=595)
    # 免责声明
    makefunc.addTesxts(fontsize=16,x=600,y=100,text="*Calculation based on 30 years tenure, 3.5% bank interest rate. For your personal")
    makefunc.addTesxts(fontsize=16,x=620,y=80,text=" financial calculation, please approach our salesperson for assistance.")
    doc.showPage()  # 保存当前画布页面
    logger.info('============>>PROGRESSIVE PAYMENT')

    # Page10 ===========================================================================
    # makefunc.background('10.jpg')
    # if userinfo['logo'] and Config.envs == 'release':
    #     ...
    #     makefunc.AddURLImages(Config.imgpath+userinfo['logo'],x=100,y=270,w=224,h=224) 
    # agentdata6 = [
    # ["NAME",':', userinfo['agentName']],
    # ["MOBILE",':', userinfo['mobile']],
    # ["EMAIL",':',userinfo['email']],
    # ["CEA",':', userinfo['regNum']]
    # ]

    # t = Table(agentdata6, style={
    # # ("FONT", (0, 0), (-1, -1), song, 35,50),
    # ("FONT", (0, 0), (-1, -1), 'arial', 35,50),
    # ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
    # ('ALIGN', (1, 0), (1, -1), 'CENTER')
    # })
    # t._argW[1] = 20
    # t.wrapOn(doc, 0, 0)
    # t.drawOn(doc, 350, 270)
    # doc.linkURL('https://api.whatsapp.com/send?phone= '+userinfo['mobile'], (50,160,50+630,230))
    # doc.showPage()  # 保存当前画布页面

    # Page11 ===========================================================================
    if fileinfo:
        for item in fileinfo:
            # print('page10')
            if item["page"] == 10:
                makefunc.AddURLImages(Config.imgpath+item['logo'],x=0,y=0,w=pagesize[0],h=pagesize[1]) 
                doc.showPage()  # 保存当前画布页面
    makefunc.background('11.jpg')
    doc.showPage()  # 保存当前画布页面
    logger.info('============>>10 page')

    # Page12 ===========================================================================
    makefunc.background('bgA.png')
    makefunc.AddImages('14right.jpg',x=pagesize[0]-550,y=160,w=430,h=536)
    # 项目单位总数
    makefunc.addTesxts(fontsize=40,x=pagesize[0]-530,y=610,text=str(prodatainfo['unitsNum']),fontname='ARIALBD')
    # 项目已释放单位数
    makefunc.addTesxts(fontsize=40,x=pagesize[0]-530,y=410,text=str(int(prodatainfo['unitsNum'])-int(prodatainfo['released'])),fontname='ARIALBD')
    # 项目已售单位数
    makefunc.addTesxts(fontsize=40,x=pagesize[0]-530,y=220,text=str(prodatainfo['sold']),fontname='ARIALBD')

    makefunc.addTesxts(fontsize=60,x=600,y=pagesize[1]-180,text="Sales Transactions",fontname='ARIALBD')

    # 单位销量统计
    unittransactions = [
        ['Date','Floor','Size(Sqft)','Price','Price(psf)'],
        ['','','','',''],
        ['','','','',''],
        ['','','','',''],
        ['','','','',''],
        ['','','','',''],
        ['','','','',''],
        ['','','','',''],
        ['','','','',''],
        ['','','','',''],
        ['','','','',''],
    ]
    for i,item in enumerate(dealInfo):
        date = '-'
        unitprice = 0
        transactionPrice = 0
        if item['transactionDate']:
            date = time.strftime('%Y-%m',time.localtime(item['transactionDate']/1000))
        if item['price']:
            unitprice = item['price']
        if item['transactionPrice']:
            transactionPrice = item['transactionPrice']
        unittransactions[i+1] =  [date,item['floor'],item['area'],makefunc.priceset(transactionPrice),makefunc.priceset(unitprice)]
        # unittransactions.append([date,item['floor'],item['area'],makefunc.priceset(transactionPrice),makefunc.priceset(unitprice),])
    
    t = Table(unittransactions,(pagesize[0]-560)/6,60, style={
    # ("FONT", (0, 0), (-1, -1), song, 22),
    ("FONT", (0, 0), (-1, -1), 'arial', 22),
    ('LINEABOVE', (0, 0), (-1, -1), 0.05 * cm, colors.whitesmoke),
    ('LINEBELOW', (0,-1), (-1,-1), 0.05 * cm, colors.whitesmoke),
    ('BACKGROUND',(0,0),(-1,0), colors.darkslateblue),
    ('TEXTCOLOR',(0,0),(-1,0), colors.white),
    ('BACKGROUND',(0,1),(-1,-1), colors.white),
    ("TEXTCOLOR", (0,1), (-1, -1), colors.black),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER')
    })
    t._argH[0] = 35
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 100, 125)
    doc.showPage()  # 保存当前画布页面

    # Page13 ===========================================================================
    makefunc.background('bgA.png')
    # New Project Sales Progress
    # Core Central Region (CCR)
    makefunc.addTesxts(fontsize=40,x=pagesize[0]-600,y=500,text="New Project Sales Progress")
    makefunc.addTesxts(fontsize=25,x=pagesize[0]-600,y=450,text="Core Central Region (CCR)")
    
    makefunc.make_drawing(CCRinfo).drawOn(doc,0,125)
    makefunc.AddImages('total sold.png',x=pagesize[0]-600,y=700,w=60,h=50)

    doc.showPage()  # 保存当前画布页面
    logger.info('============>>Core Central Region (CCR)')

    # Page14 ===========================================================================
    makefunc.background('bgA.png')
    # New Project Sales Progress
    # Core Central Region (RCR)
    makefunc.addTesxts(fontsize=40,x=pagesize[0]-600,y=500,text="New Project Sales Progress")
    makefunc.addTesxts(fontsize=25,x=pagesize[0]-600,y=450,text="Core Central Region (RCR)")
    makefunc.make_drawing(RCRinfo).drawOn(doc,0,125)
    makefunc.AddImages('total sold.png',x=pagesize[0]-600,y=700,w=60,h=50)
    doc.showPage()  # 保存当前画布页面
    logger.info('============>>Core Central Region (RCR)')

    # Page15 ===========================================================================
    makefunc.background('bgA.png')

    # New Project Sales Progress
    # Core Central Region (OCR)
    makefunc.make_drawing(OCRinfo).drawOn(doc,0,125)
    makefunc.addTesxts(fontsize=40,x=pagesize[0]-600,y=500,text="New Project Sales Progress")
    makefunc.addTesxts(fontsize=25,x=pagesize[0]-600,y=450,text="Core Central Region (OCR)")
    makefunc.AddImages('total sold.png',x=pagesize[0]-600,y=700,w=60,h=50)
    doc.showPage()  # 保存当前画布页面
    logger.info('============>>Core Central Region (OCR)')

    # Page16 ===========================================================================
    makefunc.background('16.jpg')
    doc.showPage()  # 保存当前画布页面

    # Page16 ===========================================================================
    makefunc.background('17.jpg')

    if userinfo['logo'] and Config.envs == 'release':
        ...
        makefunc.AddURLImages(Config.imgpath+userinfo['logo'],x=700,y=250,w=224,h=224) 

    agentdata3 = [
    ["NAME",':', userinfo['agentName']],
    ["MOBILE",':', userinfo['mobile']],
    ["EMAIL",':',userinfo['email']],
    ["CEA",':', userinfo['regNum']]
    ]

    t = Table(agentdata3, style={
    # ("FONT", (0, 0), (-1, -1), song, 35,50),
    ("FONT", (0, 0), (-1, -1), 'arial', 35,50),
    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
    ('ALIGN', (1, 0), (1, -1), 'CENTER')
    })
    t._argW[1] = 20
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 920, 250)

    # 定义 URL 各部分
    url_params = {}
    components = [
        "https",          # 协议 (scheme)
        Config.share_domain,# 域名 (netloc),
        prodatainfo['abbreviation']+"/"+userinfo['regNum'],      # 路径 (path)
        "",               # 参数 (params) - 较少使用
        parse.urlencode(url_params),  # 查询参数 (query)
        ""                # 片段 (fragment) - 如 #top
    ]
    
    # 跳转到 项目分享页面
    doc.linkURL(parse.urlunparse(components), (1250,130,1250+350,200))

    # 跳转到 whatsapp
    # doc.linkURL('https://api.whatsapp.com/send?phone= '+userinfo['mobile'], (1250,130,1250+350,200))
    doc.showPage()  # 保存当前画布页面
    logger.info('============>>add 17 page')

    # Page17 ===========================================================================
    makefunc.background('overpage.png')
    # makefunc.addTesxts(fontsize=40,x=pagesize[0]/4,y=800,text="DISCLAIMER")
    # logger.info(gettime.getDates())
    text_n1 = """
    While Huttons has endeavoured to ensure that the information and materials contained
    herein are accurate and up to date as at [{0}], Huttons is not responsible for any
    errors or omissions,or for the results obtained from their use or the reliance placed
    on them.All information is provided 'as is',with no guarantee of completeness,
    and accuracy. In no event will Huttons and/or salespersons thereof be liable in contract
        or in tort,to any party for any decision made or action taken in reliance on the information
        in this document or for any direct, indirect, consequential, special or similar damages.
    """.format(str(gettime.getDates()))
    data_text = [['DISCLAIMER'],[makefunc.create_body_text(text_n1,font_size=35,color=colors.white,leading=36)]]
    # makefunc.addTesxts(fontsize=40,x=pagesize[0]/5,y=700,text=)
    t = Table(data_text, style={
    ('TEXTCOLOR',(0,0),(-1,-1), colors.black),
    ("FONT", (0, 0), (-1, -1), 'arial', 40,65),
    # ('BACKGROUND',(0,0),(-1,0), colors.darkslateblue),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), 
    })
    t._argW[0] = pagesize[0]-400
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 200, 400)

    doc.showPage()  # 保存当前画布页面
    doc.save()  # 保存文件并关闭画布
    # time.sleep(3)
    # return returnPath
    return savepath 

# PND项目对比PDF生成函数
def ComparisonPDF(agentId,projectId):
    ##################################################################
    # 项目对比报表
    ##################################################################
    proidlist = projectId.split(',')
    getapi = getAPI()
    # 项目信息
    prolsit = []
    proinfourl = Config.API_IP+Config.PND_PROJECT_INFO
    for item in proidlist:
        proinfodata = {
            "projectId":item,
            "agentId":agentId
        }
        proinfo = getapi.requsetAPI(proinfourl,proinfodata)
        # prolsit[item] = proinfo
        prolsit.append(proinfo)
        if not proinfo['agentInfo']:
            return 
        logger.info('---------->项目信息获取完成 %s'%proinfo)
    Symbol = "$" #价格符
    pagesize = (1747,965) # 画布大小
    # pagesize = (A4[1],A4[0]) # 画布大小
    # PND文件夹+ 项目ID + 用户信息 + 文件名称
    # Config.filepath = os.getcwd() # 当前文件路径 服务器文件路径 ：/home/upload/broke/pnd/file/report
    Imagepath = os.path.join(Config.filepath,'file')
    # savepath = os.path.join(Imagepath,'test.pdf') 
    gettime = getDatetimes()
    tt = gettime.getDate()
    uppath = os.path.join(Config.returnpaths,'Comparison',tt)
    if not os.path.exists(uppath):
        os.makedirs(uppath)
    # filename = agentId+str(int(time.time()))
    savepath = os.path.join(uppath,str(int(time.time()))+'.pdf') 
    # returnPath = os.path.join(uppath,str(int(time.time()))+'.pdf')
    doc = canvas.Canvas(savepath,pagesize=pagesize)
    doc.setTitle("ProJect Comparison")
    makefunc = MakeReportlab(doc,Imagepath,pagesize,Symbol) # 加载方法
    
    # 注册字体    
    msyh = os.path.join(os.getcwd(), "font", "msyh.ttf")
    msyhbd = os.path.join(os.getcwd(), "font", "msyhbd.ttf")
    ARIALBD = os.path.join(os.getcwd(), "font", "ARIALBD.TTF")
    arial = os.path.join(os.getcwd(), "font", "arial.ttf")
    simsun = os.path.join(os.getcwd(), "font", "simsun.ttc")

    song = "simsun"
    pdfmetrics.registerFont(TTFont(song, simsun))
    pdfmetrics.registerFont(TTFont('ARIALBD',ARIALBD)) #注册字体
    pdfmetrics.registerFont(TTFont('arial',arial)) #注册字体
    pdfmetrics.registerFont(TTFont('msyh',msyh)) #注册字体
    pdfmetrics.registerFont(TTFont('msyhbd',msyhbd)) #注册字体
    # pdfmetrics.registerFont(TTFont('dejavu','dejavu-sans.book.ttf')) #注册字体
    logger.info('----------> 生成空文件')
    styles = getSampleStyleSheet()["Normal"]
    styles.leading = 18
    styles.fontSize = 16
    styles.alignment  = 1
    styles.fontName = 'ARIALBD'
    styles.textColor = HexColor("0x335C72")
    # page 1
    userinfo = {
        "agentName":"",
        "email":"",
        "mobile":"",
        "regNum":"",
        "id":"",
    }

    procomparison = [
        ["PROJECT",'','','','',''],
        ["DISTRICT",'','','','',''],
        ["TOTAL UNITS",'','','','',''],
        ["TOP",'','','','',''],
        ["TENURE",'','','','',''],
        ["DEVELOPER",'','','','',''],
        ["1 BR",'-','-','-','-','-'],
        ["2 BR",'-','-','-','-','-'],
        ["3 BR",'-','-','-','-','-'],
        ["4 BR",'-','-','-','-','-'],
        ["5 BR",'-','-','-','-','-'],
        [Paragraph("Nearby MRT \n within 2KM", style=styles),'','','','',''],
        ["SCHOOLS¹",'','','','',''],
        [Paragraph('DOWNLOAD \n BROCHURE', style=styles),'','','','',''],
        [Paragraph('360 \n PANORAMA', style=styles),'','','','',''],
    ]
    proimglist = []
    for keys,items in enumerate(prolsit):
        unitInfoinfo = items['unitInfo']
        projectdata = items['projectInfo']
        userinfo = items['agentInfo']
        Symbol = "$" #价格符
        if projectdata["currencySymbol"] != None:
            Symbol = projectdata["currencySymbol"]
        completionDate = ''
        if type(projectdata['completionDate']) is int:
            timeArray = time.strptime(projectdata['completionDate'], "%Y-%m")
            completionDate = time.strftime("%Y", timeArray)
        elif type(projectdata['completionDate']) is str and makefunc.isVaildDate(projectdata['completionDate']):
            a = time.strptime(projectdata['completionDate'], "%Y-%m-%d %H:%M:%S")
            completionDate = time.strftime("%Y", a)
        else:
            completionDate = projectdata['completionDate']
        # print(int(keys+1),projectdata)
        procomparison[0][int(keys+1)] = projectdata['projectName']
        procomparison[1][keys+1] = projectdata['district']
        procomparison[2][keys+1] = projectdata['unitsNum']
        procomparison[3][keys+1] = completionDate
        procomparison[4][keys+1] = projectdata['tenure']
        procomparison[5][keys+1] = projectdata['brokeName']
        proimglist.append(projectdata['mainImage'])
        for unit in unitInfoinfo:
            # print(unit)
            # 项目单位处理
            if unit['bedrooms'] == 1 and unit['price']:
                if unit['ivt']:
                    procomparison[6][keys+1] = Symbol+format(round(unit['price']),',') + '    <IVT>'
                else:
                    procomparison[6][keys+1] = Symbol+format(round(unit['price']),',')
            if unit['bedrooms'] == 2 and unit['price']:
                if unit['ivt']:
                    procomparison[7][keys+1] = Symbol+format(round(unit['price']),',') + '    <IVT>'
                else:   
                    procomparison[7][keys+1] = Symbol+format(round(unit['price']),',') 
            if unit['bedrooms'] == 3 and unit['price']:
                if unit['ivt']:
                    procomparison[8][keys+1] = Symbol+format(round(unit['price']),',') + '    <IVT>'
                else:
                    procomparison[8][keys+1] = Symbol+format(round(unit['price']),',')
            if unit['bedrooms'] == 4 and unit['price']:
                if unit['ivt']:
                    procomparison[9][keys+1] = Symbol+format(round(unit['price']),',') + '    <IVT>'
                else:
                    procomparison[9][keys+1] = Symbol+format(round(unit['price']),',')
            if unit['bedrooms'] == 5 and unit['price']:
                if unit['ivt']:
                    procomparison[10][keys+1] = Symbol+format(round(unit['price']),',') + '    <IVT>'
                else:
                    procomparison[10][keys+1] = Symbol+format(round(unit['price']),',') 
        
        school = []
        MRT = []
        if projectdata['facilitiesMap']:
            facilities = json.loads(projectdata['facilitiesMap'])
            # print(facilities)
            if facilities :
                for item in facilities:
                    if item['type'] == 'subway_station':
                        if len(item['value']) > 2:
                            for MRTs in item['value'][0:2]:
                                MRT.append(MRTs['name'])
                                # print(MRTs['name'])
                        else:
                            for MRTs in item['value']:
                                MRT.append(MRTs['name'])
                                # print(MRTs['name'])
                    if item['type'] == 'school':
                        if len(item['value']) > 5:
                            for schools in item['value'][0:5]:
                                school.append(schools['name'])
                                # print(schools['name'])
                        else:
                            for schools in item['value']:
                                school.append(schools['name'])
                                # print(schools['name'])
                        # for school in item['value']:
                        #     print(school['name'])
        procomparison[11][keys+1] =makefunc.create_body_text(' / '.join(MRT),'msyh') 
        procomparison[12][keys+1] = makefunc.create_body_text(' / '.join(school),'msyh')
        if projectdata['url'] != None:
            procomparison[13][keys+1] = 'CLICK HERE'
            # doc.linkURL(Config.imgpath+projectdata['url'], (280,185,280+150,185+22))
        if projectdata['ivtList'] != []:
            procomparison[14][keys+1] = 'CLICK HERE'
            # print(prodatainfo['ivtList'][0])
            # doc.linkURL(projectdata['ivtList'][0], (280,150,280+150,150+22))

    # page1  ===============================================
    # 背景 
    logger.info('---------->proimglist 处理完成')
    makefunc.background('comparisn0.jpg')
  
    if userinfo['logo'] and Config.envs == 'release':
        ...
        makefunc.AddURLImages(Config.imgpath+userinfo['logo'],x=720,y=250,w=224,h=224) 
    agentdata4 = [
    ["NAME",':', userinfo['agentName']],
    ["MOBILE",':', userinfo['mobile']],
    ["EMAIL",':',userinfo['email']],
    ["CEA",':', userinfo['regNum']]
    ]

    t = Table(agentdata4, style={
    # ("FONT", (0, 0), (-1, -1), song, 35,50),
    ("FONT", (0, 0), (-1, -1), 'arial', 35,50),
    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
    ('ALIGN', (1, 0), (1, -1), 'CENTER')
    })
    t._argW[1] = 20
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 950, 250)
    doc.showPage()  # 保存当前画布页面
    logger.info('---------->page 1')

    makefunc.background('bgC.png')

    itemx = 240
    for imgitem in proimglist:
        makefunc.ImageAdaptive(Config.imgpath+imgitem,x=itemx,y=pagesize[1]-255,w=200,h=100) 
        itemx+=312
    t = Table(procomparison,312,42, style={
    # ("FONT", (0, 0), (0, -1), 'msyh', 18,25),

    ("FONT", (0, 0), (-1, -1), 'msyhbd', 16,25),
    ("FONT", (0, 0), (0, -1), 'ARIALBD', 16,25),
    ("FONT", (0, 6), (10, -1), 'ARIALBD', 16,25),
    ("FONT", (1,1), (-1, 5), 'arial', 18,25),
    ("TEXTCOLOR", (0, 0), (-1, -1), HexColor("0x335C72")),
    ('LINEBEFORE', (0, 0), (-1, -1), 0.05* cm, HexColor("0x335C72")),
    ('LINEBELOW', (0, 6), (-1, 10), 0.05* cm, HexColor("0x335C72")),
    ('BACKGROUND',(0,0),(-1,0), colors.white),
    ('BACKGROUND',(0,1),(-1,1), '#E6E6E6'),
    ('BACKGROUND',(0,2),(-1,2), colors.white),
    ('BACKGROUND',(0,3),(-1,3), '#E6E6E6'),
    ('BACKGROUND',(0,4),(-1,4), colors.white),
    ('BACKGROUND',(0,5),(-1,5), '#E6E6E6'),
    ('BACKGROUND',(0,6),(-1,10), colors.white),
    ('BACKGROUND',(0,11),(-1,11), '#E6E6E6'),
    ('leading',(0,11),(-1,11), 18),
    ('BACKGROUND',(0,12),(-1,12), colors.white),
    ('BACKGROUND',(0,13),(-1,13), '#E6E6E6'),
    ('BACKGROUND',(0,14),(-1,14), colors.white),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), 
    })
    t._argH[11] = 43
    t._argH[12] = 105
    t._argH[13] = 43
    t._argH[14] = 43
    t._argW[0] = 150
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 20, 25)
    # IVT 跳转链接 
    for keys,items in enumerate(prolsit):
        if items['projectInfo']['url'] != None:
            # print(items['projectInfo']['url'])
            doc.linkURL(Config.imgpath+items['projectInfo']['url'], (170+(312*(keys)),62,170+(312*(keys))+312,62+42))
        if items['projectInfo']['ivtList'] != []:
            # print(items['ivtList'])
            doc.linkURL(items['projectInfo']['ivtList'][0], (170+(312*(keys)),20,170+(312*(keys))+312,20+42))
        for unit in items['unitInfo']:
            if unit['ivt'] != None:
                ivt = unit['ivt'].split(',')
                if unit['bedrooms'] == 1:
                    doc.linkURL(ivt[0], (170+(312*(keys)),428,170+(312*(keys))+312,428+42))
                if unit['bedrooms'] == 2:
                    doc.linkURL(ivt[0], (170+(312*(keys)),386,170+(312*(keys))+312,386+42))
                if unit['bedrooms'] == 3:
                    doc.linkURL(ivt[0], (170+(312*(keys)),344,170+(312*(keys))+312,344+42))
                if unit['bedrooms'] == 4:
                    doc.linkURL(ivt[0], (170+(312*(keys)),302,170+(312*(keys))+312,302+42))
                if unit['bedrooms'] == 5:
                    doc.linkURL(ivt[0], (170+(312*(keys)),260,170+(312*(keys))+312,260+42))
    makefunc.addTesxts(fontsize=18,x=20,y=10,color=colors.white,text="¹Source: Google Maps")
    doc.showPage()  # 保存当前画布页面
    logger.info('---------->page 2')
    # makefunc.background('18.jpg')
    makefunc.background('overpage.png')
    # makefunc.addTesxts(fontsize=40,x=pagesize[0]/4,y=800,text="DISCLAIMER")
    # logger.info(gettime.getDates())
    text_n1 = """
    While Huttons has endeavoured to ensure that the information and materials contained
    herein are accurate and up to date as at [{0}], Huttons is not responsible for any
    errors or omissions,or for the results obtained from their use or the reliance placed
    on them.All information is provided 'as is',with no guarantee of completeness,
    and accuracy. In no event will Huttons and/or salespersons thereof be liable in contract
        or in tort,to any party for any decision made or action taken in reliance on the information
        in this document or for any direct, indirect, consequential, special or similar damages.
    """.format(str(gettime.getDates()))
    data_text = [['DISCLAIMER'],[makefunc.create_body_text(text_n1,font_size=35,color=colors.white,leading=36)]]
    # makefunc.addTesxts(fontsize=40,x=pagesize[0]/5,y=700,text=)
    t = Table(data_text, style={
    ('TEXTCOLOR',(0,0),(-1,-1), colors.black),
    ("FONT", (0, 0), (-1, -1), 'arial', 40,65),
    # ('BACKGROUND',(0,0),(-1,0), colors.darkslateblue),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), 
    })
    t._argW[0] = pagesize[0]-400
    t.wrapOn(doc, 0, 0)
    t.drawOn(doc, 200, 400)

    doc.showPage()  # 保存当前画布页面
    logger.info('---------->page 3')
    doc.save()  # 保存文件并关闭画布
    # return returnPath
    return savepath


# ERA 项目可售单位报表 生成
def ERABedroomRports(agentId,brokeId,minPrice,maxPrice,projectArea,token,source):
    gettime = getDatetimes()
    # Imagepath = os.path.join(Config.filepath,'file')
    tt = gettime.getDate()
    uppath = os.path.join('/home/upload/broke/',brokeId,'pdf',tt)
    retpath = os.path.join('/upload/broke/',brokeId,'pdf',tt)  # 因 文件请求映射地址与文件存放地址层级不一致
    if not os.path.exists(uppath):
        os.makedirs(uppath)
    savepath = os.path.join(uppath,agentId+str(int(time.time()))+'.pdf') 
    retpaths = os.path.join(retpath,agentId+str(int(time.time()))+'.pdf') 
    elements = []
    pdfmetrics.registerFont(TTFont('arial','arial.ttf')) #注册字体
    pdfmetrics.registerFont(TTFont('ARIALBD','ARIALBD.TTF')) #注册字体
    pdfmetrics.registerFont(TTFont('msyh','msyh.ttf')) #注册字体

    # 基础数据准备 =====================================================
    getapi = getAPI()
    # 项目信息
    proinfourl = Config.API_IP+Config.APP_PROJECT_BEDROOM_COUNT
    proinfodata = {
        "brokeId":brokeId,
        "minPrice":minPrice,
        "maxPrice":maxPrice,
        "projectArea":projectArea,
        "agentId":agentId,
        "token":token,
        # "timestamp":str(int(time.time())),
        "source":source
    }
    proinfo = getapi.requsetAPI(proinfourl,proinfodata)
    styletext = getSampleStyleSheet()['BodyText']
    styletext.fontName = 'msyh'
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Titles',
                            fontSize=18,
                            leading = 30,
                            alignment = 1,
                            fontName = 'arial'))
    # ---------------->Title
    years = gettime.Years()
    Region = ''
    minpic =''
    maxpic = ''
    if projectArea != None:
        Region = projectArea
    if minPrice != None:
        minpic = '>$'+str(format(round(int(minPrice)),','))
    if maxPrice != None:
        maxpic = ',<$'+str(format(round(int(maxPrice)),','))
    pagetitle = '<font fontSize=22  name="ARIALBD" color="#0000">{0} </font><font fontSize=18 name="ARIALBD" color="#BA530F">{1} Projects available ({2}{3}) </font>'.format(years,Region,minpic,maxpic)
    elements.append(Paragraph(pagetitle, styles['Titles']))
    data = [
        # ["OCR",'', '', '', ''],
        ['Project Name', '1 Bedroom', '2 Bedroom', '3 Bedroom', '4 Bedroom','5 Bedroom'],
        ]
    for item in proinfo:
        # print(item)
        data.append([Paragraph(item['projectName'], style=styletext), str(item['room1']), str(item['room2']), str(item['room3']), str(item['room4']), str(item['room5'])])
    
    t = Table(data)
    t.setStyle(TableStyle([
        ('TEXTCOLOR',(0,0),(-1,-1),colors.black),#设定单元格
        # ('SPAN',(0,0),(-1,0)),#合并单元格，合并首行
        ('BACKGROUND',(0,0),(-1,0),'#BA530F'),
        ('TEXTCOLOR',(0,0),(-1,0),colors.white),
        ("FONT", (0, 0), (-1, -1), 'arial', 12,16),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),#设定居中对齐
        ('INNERGRID', (0,0), (-1,-1), 0.25, colors.black),
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ]))
    t._argW[0] = 120
    t._argH[0] = 18

    elements.append(t)
    # def AllPageSetup(canvas, doc):
    #     # 文档背景图
    #     canvas.saveState()
    #     # canvas.drawImage(os.path.join(Imagepath,'test1.jpg'),0,0,8.5 * inch,12 * inch)
    #     canvas.restoreState()
    # doc.build(elements,onFirstPage=AllPageSetup,onLaterPages=AllPageSetup)
    doc = SimpleDocTemplate(savepath,topMargin=0.5*inch,bottomMargin=0.5*inch)
    doc.build(elements)
    return retpaths

# 项目分享PDF生成函数
def Shera_to_Pdf(agentId,projectId):

    # 基础数据准备 =====================================================
    getapi = getAPI()
    datas = {} 
    datas['imgpath'] = Config.imgpath
    datas['temppath'] = Config.ecoprop_temp_path # 本地图片路径

    logger.info('get User Info ====>>>>'+agentId)
    datas['userInfo'] = getapi.requsetAPI(Config.API_IP+Config.APP_AGENT_SHARE_INFO,params={"agentId": agentId})

    # # 项目信息 
    datas['proInfo'] = getapi.requsetAPI(Config.API_IP+Config.APP_PROJECT_INFO,params={"agentId": agentId,"projectId": projectId})
    if not datas['proInfo'] or not datas['userInfo']:
        raise HTTPException(status_code=404, detail="Get Pro Or User Info Error")
    logger.info('get Pro Info ====>>>>{0}'.format(projectId))
    # Base64 转码
    if datas['proInfo']['description']:
        logger.info('description Base64====>>>>')
        datas['proInfo']['description'] = base64.b64decode(datas['proInfo']['description']).decode("utf-8")
    else:
        datas['proInfo']['description'] = ''

    # 周边设施数据处理
    if datas['proInfo']['facilitiesMap']:
        logger.info('facilitiesMap  literal_eval====>>>>')
        datas['proInfo']['facilitiesMap'] = literal_eval(datas['proInfo']['facilitiesMap'])
    else:
        datas['proInfo']['facilitiesMap'] = []
    
    # 单位价格信息
    logger.info('get unit_price_list ====>>>>')
    datas['unit_price_list'] = getapi.requsetAPI(Config.API_IP+Config.APP_UNIT_TYPE_REPORT,params={"agentId": agentId,"projectId": projectId})
    
    # 项目Floor Plans
    logger.info('get site_plan_list ====>>>>')
    site_plan_list = getapi.requsetAPI_POST(Config.API_IP+Config.APP_SITEPLAN_IMG,params={"agentId": agentId,"projectId": projectId})
    site_list = {
        "siteplan":[],
        "allbuilding":[]
    }
    for site in site_plan_list:
        if site['type'] == "siteplan" and site['img']:
            site_list['siteplan'].append(site['img'])
        if site['type'] == "allbuilding" and site['img']:
            site_list['allbuilding'].append(site['img'])
    datas['site_plan'] = site_list

    # 项目媒体文件
    logger.info('get Pro_media_list ====>>>>')
    datas['media_list'] = getapi.requsetAPI_POST(Config.API_IP+Config.APP_MEDIA_SHARE,params={"agentId": agentId,"projectId": projectId})
    
    # 项目户型图列表
    logger.info('get floor_plan_list ====>>>>')
    datas['floor_plan_list'] = getapi.requsetAPI_POST(Config.API_IP+Config.APP_FLOOR_PLANS,params={"type":"","projectId": projectId,"pageNo": 1,"pageSize":10})

    # 文件内部的所有跳转全部将跳转到指定的分享页面
    logger.info('set openlink====>>>>')
    # 定义 URL 各部分
    url_params = {}
    components = [
        "https",          # 协议 (scheme)
        Config.share_domain,# 域名 (netloc)
        datas['proInfo']['abbreviation']+"/"+datas['userInfo']['regNum'],      # 路径 (path)
        "",               # 参数 (params) - 较少使用
        parse.urlencode(url_params),  # 查询参数 (query)
        ""                # 片段 (fragment) - 如 #top
    ]
    datas['openlink'] = parse.urlunparse(components)

    # 模板文件路径
    # logger.info('Set Tmp Info ====>>>>'+agentId)
    # temp_path = os.path.join(Config.ecoprop_temp_path,'ecoprop_pro_share_temp.html').replace('\\','/')
    # 输出文件路径
    logger.info('Set return Path ====>>>>')
    gettime = getDatetimes()
    tt = gettime.getDate()
    if not datas['userInfo']['regNum'] : return {'code':'-1','msg':'error',"datas":"regNum is null"}
    new_file_name = datas['proInfo']['abbreviation']+"-"+datas['userInfo']['regNum']+"-"+tt+'.pdf'
    re_path = os.path.join(Config.ecoprop_return_path,'pro_share_pdf',tt,new_file_name).replace('\\','/')

    # 文件夹检查
    if not os.path.exists(os.path.split(re_path)[0]): 
        os.makedirs(os.path.split(re_path)[0])

    logger.info('Get Jinja2 Temp ====>>>> ecoprop_pro_share_temp.html')
    env = jinja2.Environment(loader=jinja2.FileSystemLoader(searchpath=Config.ecoprop_temp_path,encoding='utf-8'))
    template = env.get_template('ecoprop_pro_share_temp.html')

    # 模板填充参数
    options = {
        # 'page-width': '842px',
        # 'page-height': '595px',
        'page-size': ['A4'],
        'margin-top': '0mm',
        'margin-right': '0mm',
        'margin-bottom': '0mm',
        'margin-left': '0mm',
        'orientation':'Landscape', #横向
        'encoding': "UTF-8",
        'no-outline': None,

        'enable-local-file-access': None,  # 允许访问本地文件
        'javascript-delay': 2000,          # 延迟执行JavaScript，确保图片完全加载
        'images': None,                    # 启用图像加载
        'enable-javascript': None,         # 启用JavaScript

        # 日志设置
        'quiet': None,                     # 减少输出信息

        # 资源加载控制
        # 内存和超时设置
        'disable-smart-shrinking': None,   # 禁用智能缩小，保持原始尺寸比例
        'zoom': 1,                         # 缩放因子
        'load-error-handling': 'ignore',   # 忽略加载错误
        'load-media-error-handling': 'ignore',  # 忽略媒体加载错误
    }
    try:
        datas = eval(re.sub('None','\'\'',str(datas))) # 去除None值
        logger.info('插入模板的参数 ====>>>>{}'.format(datas))
        htmls = template.render(datas)
        config = pdfkit_config()

        # 清空之前文件的内容
        open(re_path, 'w').close() 
        
        # Share_Pro_Preview(re_path,htmls) # 输出预览html
        
        pdfkit.from_string(htmls,re_path,options=options,configuration=config) # 生成PDF文件
        # 确保临时文件生成成功后再覆盖目标文件
        # if os.path.exists(temp_pdf_path):
        #     shutil.move(temp_pdf_path, output_path)
    except Exception as e:
        logger.error(e)
        raise HTTPException(status_code=404, detail="PDF Error")
    logger.info('PDF ADD Over ====>>>>')
    return re_path

# 输出预览html
def Share_Pro_Preview(re_path,htmls):
    preview_path = re_path.replace('.pdf', '_preview.html')
    with open(preview_path, 'w', encoding='utf-8') as preview_file:
        preview_file.write(htmls)

# 抽离wkhtmltopdf 路径配置统一维护
def pdfkit_config():
    ...
    config = pdfkit.configuration(wkhtmltopdf=Config.wkhtml_path)
    # config = ''
    return config


# 创建 单位对比PDF文件
def Share_Unit_Pdf(agentId,unitId):

    getapi = getAPI()
    datas = {
        "imgpath":Config.imgpath,
        "temppath":Config.ecoprop_temp_path, # 本地图片路径"
        "unit_list": [], # 单位列表
        "userInfo": {}, # 用户信息
    } 
    # agentId = "e73ca86d287143709c1450012bac9e9a"
    # unitId = "7da0db34a0f34d969049a4b22a768b39,7e078477280b4abfb639ab79661c4400,70911a06985b44549380cc04839a7f14,b67eaef41f224acb919852e9b98ca81a"
    logger.info('get User Info ====>>>>'+agentId)
    datas['userInfo'] = getapi.requsetAPI(Config.API_IP+Config.APP_AGENT_SHARE_INFO,params={"agentId": agentId})

    UnitIdList = unitId.split(',')
    logger.info('get Unit Info ====>>>>')
    for i in UnitIdList:
        unitInfo = getapi.requsetAPI(Config.API_IP+Config.APP_UNIT_INFO,params={"agentId": agentId,"unitId":i})
        if unitInfo : datas['unit_list'].append(unitInfo)
    if not datas['unit_list'] :
        logger.error('Unit List Is None')
        raise HTTPException(status_code=404, detail="Unit Get Error")
    # 模板文件路径
    # temp_path = os.path.join(Config.ecoprop_temp_path,'ecoprop_unit_share_temp.html').replace('\\','/')
    # 输出文件路径
    # str(datetime.datetime.now().strftime('%d-%m-%Y-%H%M%S'))
    gettime = getDatetimes()
    tt = gettime.getDate()
    new_file_name = datas['userInfo']['regNum']+"-"+tt+'.pdf'
    re_path = os.path.join(Config.ecoprop_return_path,'user_unit_compare',tt,new_file_name).replace('\\','/')
    if not os.path.exists(os.path.split(re_path)[0]):
        logger.info('ADD New folder ====>>>>'+os.path.split(re_path)[0])
        os.makedirs(os.path.split(re_path)[0])
        
    env = jinja2.Environment(loader=jinja2.FileSystemLoader(searchpath=Config.ecoprop_temp_path,encoding='utf-8'))
    template = env.get_template('ecoprop_unit_share_temp.html')
    # 模板填充参数
    options = {
        # 'page-width': '842px',
        # 'page-height': '595px',
        'page-size': ['A4'],
        'margin-top': '0mm',
        'margin-right': '0mm',
        'margin-bottom': '0mm',
        'margin-left': '0mm',
        'orientation':'Landscape', #横向
        'encoding': "UTF-8",
        'no-outline': None,

        'enable-local-file-access': None,  # 允许访问本地文件
        'javascript-delay': 2000,          # 延迟执行JavaScript，确保图片完全加载
        'images': None,                    # 启用图像加载
        'enable-javascript': None,         # 启用JavaScript

        # 日志设置
        'quiet': None,                     # 减少输出信息

        # 资源加载控制
        # 内存和超时设置
        'disable-smart-shrinking': None,   # 禁用智能缩小，保持原始尺寸比例
        'zoom': 1,                         # 缩放因子
        'load-error-handling': 'ignore',   # 忽略加载错误
        'load-media-error-handling': 'ignore',  # 忽略媒体加载错误
    }
    try:
        logger.info('Set tmp Info ====>>>>{0}'.format(datas))
        datas = eval(re.sub('None','\'\'',str(datas))) # 去除None值
        htmls = template.render(datas)
        config = pdfkit_config()
        open(re_path, 'w').close()
        pdfkit.from_string(htmls,re_path,options=options,configuration=config)
    except Exception as e:
        logger.error("File Padding Err ===>>>{0}".format(e))
        raise HTTPException(status_code=404, detail="PDF ADD Error")
        
    logger.info('PDF ADD Over ====>>>>')
    return re_path


# Ecoprop 创建 项目对比PDF文件
def Share_Pro_compare_Pdf(agentId,projectId):
    getapi = getAPI()
    datas = {
        "imgpath":Config.imgpath,
        "temppath":Config.ecoprop_temp_path, # 本地图片路径
        "pro_list": [],
    } 
    # agentId = "e73ca86d287143709c1450012bac9e9a"
    # projectId = "5b2216e95ef446bf853113450a0642f1,0c3cef5e77f547a99d61d6a2ccd37885"
    logger.info('get User Info ====>>>>'+agentId)
    datas['userInfo'] = getapi.requsetAPI(Config.API_IP+Config.APP_AGENT_SHARE_INFO,params={"agentId": agentId})

    UnitIdList = projectId.split(',')
    logger.info('get Pro Info ====>>>>')
    proinfourl = Config.API_IP+Config.APP_PROJECT_INFO
    for item in UnitIdList:
        proinfo = getapi.requsetAPI(proinfourl,{"projectId":item,"agentId":agentId})
        if proinfo: datas['pro_list'].append(proinfo) 
        
    if not datas['pro_list'] :
        logger.error('Pro List Is None')
        raise HTTPException(status_code=404, detail="Pro Get Error")

    # 模板文件路径
    # temp_path = os.path.join(Config.ecoprop_temp_path,'ecoprop_pro_compare_share_temp.html').replace('\\','/')
    # 输出文件路径
    # str(datetime.datetime.now().strftime('%d-%m-%Y-%H%M%S'))
    gettime = getDatetimes()
    tt = gettime.getDate()
    new_file_name = datas['userInfo']['regNum']+"-"+tt+'.pdf'
    re_path = os.path.join(Config.ecoprop_return_path,'user_pro_compare',tt,new_file_name).replace('\\','/')
    if not os.path.exists(os.path.split(re_path)[0]):
        logger.info('ADD New folder ====>>>>'+os.path.split(re_path)[0])
        os.makedirs(os.path.split(re_path)[0])
        
    env = jinja2.Environment(loader=jinja2.FileSystemLoader(searchpath=Config.ecoprop_temp_path,encoding='utf-8'))
    template = env.get_template('ecoprop_pro_compare_share_temp.html')

    # 定义 URL 各部分
    url_params = {
        "projectIds": projectId,  # 项目ID列表
        "agentId": agentId         # 代理人ID
    }
    components = [
        "https",          # 协议 (scheme)
        Config.share_domain,# 域名 (netloc)
        "/vsProject",      # 路径 (path)
        "",               # 参数 (params) - 较少使用
        parse.urlencode(url_params),  # 查询参数 (query)
        ""                # 片段 (fragment) - 如 #top
    ]
    datas['openlink'] = parse.urlunparse(components)


    # 模板填充参数
    options = {
        # 'page-width': '842px',
        # 'page-height': '595px',
        'page-size': ['A4'],
        'margin-top': '0mm',
        'margin-right': '0mm',
        'margin-bottom': '0mm',
        'margin-left': '0mm',
        'orientation':'Landscape', #横向
        'encoding': "UTF-8",
        'no-outline': None,

        'enable-local-file-access': None,  # 允许访问本地文件
        'javascript-delay': 2000,          # 延迟执行JavaScript，确保图片完全加载
        'images': None,                    # 启用图像加载
        'enable-javascript': None,         # 启用JavaScript

        # 日志设置
        'quiet': None,                     # 减少输出信息

        # 资源加载控制
        # 内存和超时设置
        'disable-smart-shrinking': None,   # 禁用智能缩小，保持原始尺寸比例
        'zoom': 1,                         # 缩放因子
        'load-error-handling': 'ignore',   # 忽略加载错误
        'load-media-error-handling': 'ignore',  # 忽略媒体加载错误
    }
    try:
        logger.info('Set tmp Info ====>>>>{0}'.format(datas))
        datas = eval(re.sub('None','\'\'',str(datas))) # 去除None值
        htmls = template.render(datas)
        config = pdfkit_config()
        open(re_path, 'w').close()
        pdfkit.from_string(htmls,re_path,options=options,configuration=config)
    except Exception as e:
        logger.error("File Padding Err ===>>>{0}".format(e))
        raise HTTPException(status_code=404, detail="PDF ADD Error")
        
    logger.info('PDF ADD Over ====>>>>')
    return re_path


# Word文档字段填充函数
def text_to_word(paragraph_or_cell, conntent, types):
    """Word文档字段填充"""
    
    def replace_placeholder(match_str, key):
        """替换占位符并记录日志"""
        # 去除可能的空白字符
        clean_key = key.strip()
        if clean_key in conntent:
            val = conntent[clean_key]
            # 处理 None 值
            if val is None:
                val = ""
            logger.info(f'替换: {match_str} -> {val}')
            return str(val)
        else:
            logger.warning(f'未找到内容键: {clean_key} (原匹配: {match_str})')
            return "" # 或者返回原字符串 match_str 以保留标签

    # 确定要处理的段落列表
    paragraphs_to_process = []
    if types == 'cell':
        # 如果是单元格，获取其内部所有段落
        # 注意：python-docx 中 cell.paragraphs 至少包含一个空段落
        paragraphs_to_process = paragraph_or_cell.paragraphs
    else:
        # 如果是普通段落
        paragraphs_to_process = [paragraph_or_cell]

    for paragraph in paragraphs_to_process:
        original_text = paragraph.text
        
        # 快速检查是否包含需要处理的标记
        if not any(tag in original_text for tag in ["{{", "{{img", "{{buyerList"]):
            continue

        logger.info(f"处理段落/单元格部分: '{original_text}'")

        # 特殊处理：买家列表
        if "{{buyerList}}" in original_text:
            logger.info("填充买家列表")
            paragraph.clear()
            if 'buyerList' in conntent and isinstance(conntent['buyerList'], list):
                for index, rows in enumerate(conntent['buyerList']):
                    # 避免最后一行多余换行，可根据需求调整
                    line_text = "{0} {1}".format(rows.get('num', ''), rows.get('buyerName', ''))
                    p_run = paragraph.add_run(line_text)
                    p_run.font.color.rgb = RGBColor(0, 0, 0)
                    if index < len(conntent['buyerList']) - 1:
                        paragraph.add_run('\n')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
            continue

        # 特殊处理：图片
        if "{{img" in original_text:
            logger.info("填充图片")
            matches = re.findall(r'\{\{(\w+)\}\}', original_text)
            if matches:
                img_key = matches[0]
                img_url_relative = conntent.get(img_key)
                
                paragraph.clear() # 清空原有文本
                
                if img_url_relative:
                    parsed_url = parse.urlparse(Config.imgpath + img_url_relative)
                    img_filename = os.path.basename(parsed_url.path)
                    # 防止文件名冲突或非法字符，简单处理
                    img_filename = re.sub(r'[^\w\.]', '_', img_filename)
                    img_path = os.path.join(Config.ecoprop_temp_path, img_filename).replace('\\', '/')
                    
                    try:
                        # 下载图片
                        full_url = Config.imgpath + img_url_relative
                        response = requests.get(full_url, timeout=10)
                        if response.status_code == 200:
                            with open(img_path, 'wb') as file:
                                file.write(response.content)
                            
                            # 添加图片
                            run = paragraph.add_run()
                            # 设置图片大小，可以根据需要动态调整
                            run.add_picture(img_path, width=Inches(5.5))
                            
                            # 删除临时文件
                            if os.path.exists(img_path):
                                os.remove(img_path)
                        else:
                            logger.error(f"图片下载失败: {full_url}, Status: {response.status_code}")
                            paragraph.add_run("[Image Load Failed]")
                    except Exception as e:
                        logger.error(f"图片处理异常: {e}")
                        paragraph.add_run("[Image Error]")
                else:
                    logger.warning(f"图片键 {img_key} 在数据中未找到")
            continue

        # 通用文本替换逻辑
        # 使用正则分割，保留分隔符
        # 注意：re.split 如果捕获组匹配，分隔符也会包含在结果列表中
        parts = re.split(r'(\{\{.*?\}\})', original_text)
        
        paragraph.clear() # 清空当前段落内容
        
        for part in parts:
            if not part:
                continue
                
            # 检查是否是占位符
            placeholder_match = re.match(r'\{\{(.*?)\}\}', part)
            
            if placeholder_match:
                key = placeholder_match.group(1).strip()
                
                # 处理特殊标签 (Sign, ck, rd)
                if key.startswith(("Sign", "ck", "rd")):
                    logger.info(f'特殊标签跳过: {part}')
                    # 如果希望保留标签但不可见，可以取消下面注释
                    run_white = paragraph.add_run(part)
                    run_white.font.color.rgb = RGBColor(255, 255, 255)
                    continue
                
                # 普通替换
                replacement_text = replace_placeholder(part, key)
                run_black = paragraph.add_run(replacement_text)
                run_black.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # 普通文本
                run_black = paragraph.add_run(part)
                run_black.font.color.rgb = RGBColor(0, 0, 0)

# 生成PDI签名文件,java 那边处理不了替换的问题。
@router.post("/Make_Signature_File")
def Make_Signature_Files(p: Person):
    # 填充模板参数内容
    logger.info("Make_Signature_Files ====>>>>>{0}".format(p))
    try:
        # 读取模板文件
        doc = Document(p.docPath)
        for paragraph in doc.paragraphs:
            # 遍历文档段落
            if '{{' in paragraph.text:
                text_to_word(paragraph, p.data,'text')
                
        for table in doc.tables:
            # 遍历文档中的每个表格
            for row in table.rows:
                for cell in row.cells:
                    if '{{' in cell.text:
                        text_to_word(cell, p.data,'cell')
        
        # 保存渲染后的文件
        if not p.filepath :
            p.filepath = os.path.join(Config.ecoprop_return_path,'SignatureFlie',f"{int(time.time())}.docx").replace('\\','/')
        
        # 文件夹不存在就创建
        parent_dir = os.path.dirname(p.filepath)
        os.makedirs(parent_dir, exist_ok=True)
        
        doc.save(p.filepath)
        logger.info("WORD 生成完成 ====>>>>>{0}".format(p.filepath))

        # # Word 转 PDF ==》 临时解决方案，java 那边的转换工具使用有问题
        # pdf_file = convert_word_to_pdf(p.filepath)
        # # 返回结果调整
        # return { 'code':'0', 'msg':'succeed', "datas":pdf_file } 
    
        return { 'code':'0', 'msg':'succeed', "datas":p.filepath }
    
    except BaseException as e:
        logger.info("WORD 生成失败 ====>>>>>{0}".format(e))
        return {'code':'-1','msg':'error',"datas":e}




def convert_word_to_pdf(input_path, output_dir=None):
    """
    使用LibreOffice将Word文档转换为PDF
    :param input_path: Word文件路径（.docx或.doc）
    :param output_dir: 输出PDF目录，默认与输入文件同目录
    :return: 转换后的PDF路径，失败则返回None
    """
    # 检查输入文件是否存在
    if not os.path.exists(input_path):
        logger.info(f"错误：文件不存在 - {input_path}")
        return None
    
    # 确定输出目录
    if output_dir is None:
        output_dir = os.path.dirname(input_path)
    os.makedirs(output_dir, exist_ok=True)  # 确保输出目录存在
    
    # 构建输出PDF路径（与输入文件同名）
    file_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = os.path.join(output_dir, f"{file_name}.pdf")
    
    # 构建LibreOffice命令
    # 使用--headless无界面模式，--convert-to指定格式，--outdir指定输出目录
    cmd = [
        "soffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        input_path
    ]
    
    try:
        # 执行命令
        result = subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        logger.info(f"转换成功：{output_path}")
        return output_path
    except subprocess.CalledProcessError as e:
        logger.error(f"转换失败：{e.stderr}")
        return None