
import re,os
from docx import Document
from docx.shared import Inches,RGBColor
from docx.text.paragraph import Paragraph
from urllib import parse
from abc import ABC, abstractmethod
from comm.logger import logger
from typing import Dict, List, Optional,Callable,Any
from PIL import Image
from comm.logger import logger
from config import Config



class TemplateLoopProcessor:
    """支持列表循环填充和缺失字段处理的模板处理器"""
    
    def __init__(self, 
                 placeholder_pattern: str = r'{{\s*(\w+)\s*}}',
                 loop_start_pattern: str = r'{%\s*loop\s+(\w+)\s+in\s+(\w+)\s*%}',
                 loop_end_pattern: str = r'{%\s*endloop\s*%}',
                 image_placeholder: str = r'{{\s*image:\s*(\w+)\s*}}'):
        """初始化处理器"""
        self.placeholder_pattern = re.compile(placeholder_pattern)
        self.image_placeholder = re.compile(image_placeholder) # 图片占位符: {{image:field}}
        self.loop_start_pattern = re.compile(loop_start_pattern)
        self.loop_end_pattern = re.compile(loop_end_pattern)
        self.missing_handler = None  # 缺失字段处理器
    
    def set_missing_handler(self, handler: Callable[[str, str], str]):
        """设置缺失字段处理器"""
        self.missing_handler = handler
    
    def _default_missing_handler(self, field: str, original: str) -> str:
        """默认缺失字段处理：保留原始占位符"""
        return ""
        # return original
    
    def _process_placeholders(self, text: str, data: Dict[str, Any], 
                             exclude_fields: List[str] = None) -> str:
        """处理普通占位符替换，包含缺失字段处理"""
        exclude_fields = exclude_fields or []
        # print(f"处理普通占位符替换{text}")
        def replace_placeholder(match):
            field = match.group(1)
            original_placeholder = match.group(0)  # 原始占位符（如{{field}}）
            
            # 处理排除字段
            if field in exclude_fields:
                return original_placeholder
            
            # 处理存在的字段
            if field in data:
                return str(data[field])
            
            # 处理缺失的字段
            handler = self.missing_handler or self._default_missing_handler
            return handler(field, original_placeholder)
        return self.placeholder_pattern.sub(replace_placeholder, text)
    def _insert_image(self, paragraph, image_path: str, width: Optional[float] = None) -> None:
        """在指定段落插入图片"""
        if not os.path.isfile(image_path):
            raise FileNotFoundError(f"图片文件不存在: {image_path}")
        
        # 清除段落中的占位符文本
        # paragraph.text = paragraph.text.replace(f"{{image:{os.path.basename(image_path)}}}", "")
        # 图片填充段落内容清除，当前填充是全屏填充效果
        paragraph.text = ""
        run = paragraph.add_run()

        # 使用 Pillow 获取图片的原始宽高
        with Image.open(image_path) as img:
            original_width, original_height = img.size

        # 如果指定了宽度，按比例计算高度
        if width:
            aspect_ratio = original_height / original_width
            picture_width = Inches(width)
            picture_height = Inches(width * aspect_ratio)
            picture = run.add_picture(image_path, width=picture_width, height=picture_height)

        else:
            # 如果未指定宽度，直接插入图片
            picture = run.add_picture(image_path)
        # 图片居中对齐
        # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _process_loop(self, text: str, data: Dict[str, Any], 
                     exclude_fields: List[str] = []) -> str:
        """处理循环块，支持嵌套循环"""
        # 查找循环开始标记
        loop_start = self.loop_start_pattern.search(text)
        if not loop_start:
            return text  # 没有循环块，直接返回
        
        # 提取循环参数：{% loop item in list %} → item_var=item, list_name=list
        item_var, list_name = loop_start.groups()
        
        # 查找对应的循环结束标记
        loop_end = self.loop_end_pattern.search(text, pos=loop_start.end())
        if not loop_end:
            return text  # 没有匹配的结束标记，视为无效循环
        
        # 提取循环块内容
        loop_block_start = loop_start.end()
        loop_block_end = loop_end.start()
        loop_content = text[loop_block_start:loop_block_end]
        
        # 获取循环数据（如果不存在则返回空列表）
        loop_data = data.get(list_name, [])
        if not isinstance(loop_data, list):
            loop_data = []
        
        # 处理循环内容：为列表中的每个项生成填充后的内容
        processed_items = []
        for item_data in loop_data:
            # 如果列表项不是字典，包装成字典以便统一处理
            if not isinstance(item_data, dict):
                item_data = {item_var: item_data}
            # print(f'处理循环填充：{loop_content}')
            # 替换当前循环项的占位符（包含缺失字段处理）
            processed_item = self._process_placeholders(
                loop_content, 
                item_data, 
                exclude_fields
            )
            processed_items.append(processed_item  + "\n")
        
        # 拼接循环结果：用处理后的内容替换整个循环块
        processed_loop = ''.join(processed_items)
        new_text = text[:loop_start.start()] + processed_loop + text[loop_end.end():]
        
        # 递归处理可能存在的嵌套循环
        return self._process_loop(new_text, data, exclude_fields)

    # 处理白色文本填充
    def _process_white_text(self, paragraph, data: Dict[str, Any], 
                     exclude_fields: List[str] = []) :
        
        # 定义目标标签常量
        TARGET_TAGS = ("{{Sign", "{{ck", "{{rd")
        if not any(tag in paragraph.text for tag in TARGET_TAGS):
            return ''
        def clear_cell(cell):
            # 遍历单元格内的所有段落
            for para in cell.paragraphs:
                # 清空段落内容（保留段落对象）
                para.clear()

        def is_table_paragraph(paragraph) -> bool:
            """检查段落的父级是否为表格单元格"""
            return paragraph._parent.__class__.__name__ == "_Cell"
        
    
        logger.info(f"需要转白色的文本段落 =====>{paragraph.text}")
        # 使用正则表达式拆分段落内容
        parts = re.split(r'(\{\{.*?\}\})', paragraph.text)
        # 清空原有段落内容
        if is_table_paragraph(paragraph):  # 判断是否为表格内的段落
            clear_cell(paragraph)  # 清除表格内容
            current_para = paragraph.paragraphs[0]  # 关键：获取段落对象
        else:
            # 普通段落直接清空
            paragraph.clear()
            current_para = paragraph  # 当前段落就是操作对象
        
        # 遍历拆分后的部分，分别处理
        for part in parts:
            if not part:  # 跳过空字符串
                continue
            
            # 判断内容开头
            if part.startswith(("{{Sign", "{{ck", "{{rd")):
                run = current_para.add_run(part)  # 正确：用段落对象调用add_run
                run.font.color.rgb = RGBColor(255, 255, 255)  # 白色
            # 其他内容保持默认样式
            else:
                current_para.add_run(part)  # 同样用段落对象

    def process(self, paragraph, data: Dict[str, Any], 
               exclude_fields: Optional[List[str]] = [],
               missing_handler: Optional[Callable[[str, str], str]] = None) -> str:
        """
        处理整个模板
        :param paragraph.text: 模板文本
        :param data: 填充数据
        :param exclude_fields: 需要排除的字段
        :param missing_handler: 缺失字段处理器（优先级高于实例设置的处理器）
        :return: 处理后的文本
        """
        original_text = paragraph.text
        # 先处理图片
        if self.image_placeholder.search(original_text):
            # logger.info(f'开始图片填充 ====>>>>{original_text}')

            # 优先处理图片占位符
            for match in self.image_placeholder.finditer(original_text):
                field = match.group(1)
                original = match.group(0)
                
                # 是否是排除字段
                if field in exclude_fields:
                    continue

                # 填充参数是否包含图片key
                if field in data:
                    try:
                        self._insert_image(paragraph, data[field], width=6)  # 图片宽度4英寸
                    except Exception as e:
                        handler = self.missing_handler or self._default_missing_handler
                        paragraph.text  = original_text.replace(original, handler(field, f"[图片错误: {str(e)}]"))
                else:
                    handler = self.missing_handler or self._default_missing_handler
                    paragraph.text  = original_text.replace(original, handler(field, original))
                return
        
        # 普通填充文本
        text = ''.join(run.text for run in paragraph.runs) # 处理段落内容识别断行问题
        # logger.info(f'普通填充文本填充 ====>>>>{text}')


        # 临时设置缺失字段处理器（参数优先级更高）
        if missing_handler:
            original_handler = self.missing_handler
            self.missing_handler = missing_handler
        
        # 先处理所有循环块
        processed_text = self._process_loop(text, data, exclude_fields)
        # 再处理剩余的普通占位符
        paragraph.text  = self._process_placeholders(processed_text, data, exclude_fields)
        # logger.info(f'内容填充完成 ====>>>>{paragraph.text}')
        # 恢复原始处理器
        if missing_handler:
            self.missing_handler = original_handler
        
        return processed_text

class TemplateFiller(ABC):
    """模板填充器基类"""
    
    def __init__(self, template_path: str):
        """
        初始化模板填充器
        :param template_path: 模板文件路径
        """
        self.template_path = template_path
        self.data = {}
        self.excluded_fields = []
    
    @abstractmethod
    def load_template(self) -> None:
        """加载模板"""
        pass
    
    @abstractmethod
    def fill(self, data: Dict[str, str], excluded_fields: Optional[List[str]] = None) -> None:
        """
        填充模板
        :param data: 填充数据字典
        :param excluded_fields: 排除不填充的字段列表
        """
        self.data = data
        self.excluded_fields = excluded_fields if excluded_fields else []
    
    @abstractmethod
    def save(self, output_path: str) -> None:
        """
        保存填充后的文件
        :param output_path: 输出文件路径
        """
        logger.info(f'保存填充后的文件 ====>>>>{output_path}')

        pass



    @abstractmethod
    def save_pdf(self, output_path: str) -> None:
        """
        保存填充后的Word文档为PDF格式
        """
        logger.info(f'保存填充后的文件，并转为PDF ====>>>>{output_path}')
        pass

class WordTemplateFiller(TemplateFiller):
    """Word模板填充器"""
    
    def __init__(self, template_path: str):
        super().__init__(template_path)
        self.document = None
        self.load_template()
    
    def load_template(self) -> None:
        """加载Word模板"""
        self.document = Document(self.template_path)
    
    def fill(self, data: Dict[str, str], excluded_fields: Optional[List[str]] = None) -> None:
        """
        填充Word模板
        模板中的占位符格式为 {{field_name}}
        """

        logger.info(f'开始处理 ====>>>>')
        super().fill(data, excluded_fields)
        
        if not self.document:
            raise Exception("模板未加载，请先调用load_template方法")
        
        # 遍历文档中的所有段落
        for paragraph in self.document.paragraphs:
            self._fill_paragraph(paragraph)
        
        # 遍历文档中的所有表格
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._fill_paragraph(paragraph)
    
    def _fill_paragraph(self, paragraph):
        """填充单个段落"""
        replacer = TemplateLoopProcessor()
        # 区分段落是否包含图片填充
        replacer.process(paragraph, self.data, self.excluded_fields)
        # print(f"填充结果 =====>{paragraph.text}")
        
    def save(self, output_path: str) -> None:
        """保存填充后的Word文档"""
        if not self.document:
            raise Exception("模板未加载或未填充数据")
        
        self.document.save(output_path)
    def save_pdf(self, output_path: str) -> None:
        """
        保存填充后的Word文档为PDF格式
        """

        if not self.document:
            raise Exception("模板未加载或未填充数据")
        
        # 首先保存为临时的 Word 文件
        temp_word_path = output_path.replace('.pdf', '_temp1.docx')
        
        try:
            self.document.save(temp_word_path) # 保存为临时 Word 文件
            # 使用 docx2pdf 转换为 PDF
            from docx2pdf import convert
            convert(temp_word_path, output_path)
        except ImportError:
            raise Exception("请安装 docx2pdf 库: pip install docx2pdf")
        except Exception as e:
            logger.error(e)
            raise Exception(f"转换为PDF失败: {str(e)}")
        finally:
            # 删除临时文件
            import os
            if os.path.exists(temp_word_path):
                logger.info(f'删除临时文件 ====>>>>{temp_word_path}')
                os.remove(temp_word_path)

class HtmlTemplateFiller(TemplateFiller):
    """HTML模板填充器"""
    
    def __init__(self, template_path: str, use_jinja2: bool = False):
        super().__init__(template_path)
        self.html_content = ""
        self.use_jinja2 = use_jinja2
        self.load_template()
        logger.info(f'加载模板 ====>>>>{template_path}')
    def load_template(self) -> None:
        """加载HTML模板"""
        with open(self.template_path, 'r', encoding='utf-8') as f:
            self.html_content = f.read()
    
    def fill(self, data: Dict[str, str], excluded_fields: Optional[List[str]] = None) -> None:
        """
        填充HTML模板
        模板中的占位符格式为 {{field_name}}
        """
        super().fill(data, excluded_fields)
        
        if self.use_jinja2:
            self._fill_with_jinja2()
        else:
            self._fill_with_simple_replace()
    
    def _fill_with_simple_replace(self):
        """使用简单字符串替换填充"""
        for field, value in self.data.items():
            if field in self.excluded_fields:
                continue
            self.html_content = self.html_content.replace(f'{{{{{field}}}}}', str(value))
    
    def _fill_with_jinja2(self,tmplate_path: str = 'ecoprop_pro_share_temp.html'):
        """使用Jinja2模板引擎填充"""
        try:
            from jinja2 import Template
        except ImportError:
            raise ImportError("请安装jinja2库以使用此功能: pip install jinja2")
        
        # 过滤掉排除的字段
        filtered_data = {k: v for k, v in self.data.items() if k not in self.excluded_fields}

        filtered_data = eval(re.sub('None','\'\'',str(filtered_data))) # 去除None值
        template = Template(self.html_content)
        self.html_content = template.render(**filtered_data)
    
    def save(self, output_path: str) -> None:
        """保存填充后的HTML文档"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(self.html_content)
    
    def save_pdf(self, output_path: str) -> None:
        """保存填充后的HTML文档为PDF"""
        logger.info(f'保存填充后的HTML文档为PDF ====>>>>{output_path}')
        try:
            import pdfkit
        except ImportError: 
            raise ImportError('请安装pdfkit模块')

        options = {
            # 'page-width': '842px',
            # 'page-height': '595px',
            'page-size': ['A4'],
            'margin-top': '0mm',
            'margin-right': '0mm',
            'margin-bottom': '0mm',
            'margin-left': '0mm',
            # 'orientation':'Landscape', #横向
            'orientation':'portrait', #纵向
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': None,  # 允许访问本地文件
            'quiet': None
        }
        config = pdfkit.configuration(wkhtmltopdf=Config.wkhtml_path) # 当环境提示找不到wkhtmltopdf时，请取消注释并填写wkhtmltopdf的绝对路径
        
        # 清空目标文件内容，否则会出现文件内容越来越大（具体原因没有深度探索）
        with open(output_path, 'w') as f:
            f.truncate(0)

        pdfkit.from_string(self.html_content,output_path,options=options,configuration=config)
        # pdfkit.from_string(self.html_content,output_path,options=options)




# 使用示例
# if __name__ == "__main__":
#     # 示例数据
    # data = {
    #     "signDate": "张三",
    #     "buyerList": "30",
    #     "buyerName1": "北京市朝阳区",
    #     "text": "13800138000",
    #     "lawyerName": "zhangsan@example.com",
    #     "companyName1":"666",
    #     "order_id": "ORD-2023-5678",
    #     "customer_name": "王五",
    #     "total_amount": 8997,
    #     "cover_image": R"C:\Users\ycc84\Downloads\Elevation_Chart.png",  # 替换为你的图片路径
    #     "products": [
    #         {"name": "笔记本电脑", "price": 5999, "quantity": 1, "subtotal": 5999,"cover_image": R"C:\Users\ycc84\Downloads\Elevation_Chart.png",},
    #         {"name": "无线鼠标", "price": 199, "quantity": 2, "subtotal": 398,"cover_image": R"C:\Users\ycc84\Downloads\Elevation_Chart.png",},
    #         {"name": "键盘", "price": 499,"subtotal": 499},
    #         {"name": "显示器", "price": 2499, "quantity": 1, "subtotal": 2499,"cover_image": R"C:\Users\ycc84\Downloads\Elevation_Chart.png",}
    #     ]
    # }
    
#     # 要排除的字段
#     excluded_fields = ["Sign Buyer1","buyerList"]

    # Word模板填充示例
    # try:
    #     word_filler = WordTemplateFiller(R"C:\Users\ycc84\Downloads\PDI or OPT tmp - 副本.docx")
    #     word_filler.fill(data, excluded_fields)
    #     word_filler.save("filled_word.docx")
    #     print("Word模板填充完成")
    # except Exception as e:
    #     print(f"Word模板处理出错: {str(e)}")
    
    # # HTML模板填充示例（简单替换）
    # try:
    #     html_filler = HtmlTemplateFiller(R"D:\work\pythonJB\py_jb\testPY\Template.html")
    #     html_filler.fill(data, excluded_fields)
    #     html_filler.save("filled_html_simple.html")
    #     print("HTML模板(简单替换)填充完成")
    # except Exception as e:
    #     print(f"HTML模板(简单替换)处理出错: {str(e)}")
    
    # HTML模板填充示例（Jinja2）
    # try:
    #     html_jinja_filler = HtmlTemplateFiller(R"D:\work\pythonJB\py_jb\testPY\Template.html", use_jinja2=True)
    #     html_jinja_filler.fill(data, excluded_fields)
    #     html_jinja_filler.save("filled_html_jinja.html")
    #     print("HTML模板(Jinja2)填充完成")
    # except Exception as e:
    #     print(f"HTML模板(Jinja2)处理出错: {str(e)}")

